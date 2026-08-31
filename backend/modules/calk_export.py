"""
modules/calk_export.py
=======================
Export CALK (Catatan atas Laporan Keuangan) dwibahasa Indonesia/Inggris
dalam format Word (.docx) + PDF, mengikuti gaya template referensi yang
diberikan user (2 kolom bahasa berdampingan, angka menghadap tengah,
label menghadap tepi luar; kolom Indonesia = periode berjalan, kolom
Inggris = periode pembanding -- lihat catatan di susun_calk_docx()).

SUMBER DATA (dipakai ulang, TIDAK menghitung ulang apa pun):
  - `neraca`    : output susun_neraca() di laporan_keuangan.py
                  (dict berisi "aset"/"liabilitas"/"ekuitas", tiap item
                  akun punya no_akun/nama_akun/sub_kategori/saldo_akhir/
                  saldo_awal/normal_saldo -- lihat hitung_saldo_per_akun())
  - `laba_rugi` : output susun_laba_rugi() (dict "pendapatan"/"beban")
  - `coa`       : list akun COA client (dipakai fallback nama & sub_kategori)
  - `profil`    : dict profil perusahaan utk Note 1 "Umum" (lihat
                  ProfilPerusahaan di bawah) -- BELUM ada tabelnya di
                  database (Client model cuma nama/lokasi/kontak), jadi
                  wajib disuplai pemanggil (mis. lewat form khusus di
                  frontend) sampai kolom2 ini ditambahkan ke DB.

CARA PAKAI SINGKAT:
    from modules.calk_export import export_calk

    hasil = export_calk(
        output_dir="/path/output",
        nama_file_dasar="CALK_PT_AADL_2026",
        nama_perusahaan="PT AKSHAYA ANANDA DYLAN LESTARI",
        tanggal_neraca=date(2026, 7, 31),
        tanggal_neraca_lalu=date(2025, 12, 31),
        neraca=neraca,               # dari susun_neraca()
        laba_rugi=laba_rugi,         # dari susun_laba_rugi()
        coa=coa,
        profil=profil_perusahaan,    # dict, lihat contoh CONTOH_PROFIL di bawah
    )
    # hasil == {"docx": "...docx", "pdf": "...pdf"}
"""
from __future__ import annotations

import os
import re
import json
import locale
import subprocess
from datetime import date, datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .logging_config import get_module_logger
from .calk_mapping import kelompokkan_akun_calk
from .claude_client import panggil_claude_teks, ClaudeError

logger = get_module_logger("calk_export")

_FONT_UTAMA = "Times New Roman"
_UKURAN_NORMAL = 10.5
_UKURAN_JUDUL = 13
_WARNA_HITAM = RGBColor(0, 0, 0)

_BULAN_ID = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli",
             "Agustus", "September", "Oktober", "November", "Desember"]
_BULAN_EN = ["January", "February", "March", "April", "May", "June", "July",
             "August", "September", "October", "November", "December"]
_BULAN_EN_SINGKAT = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul",
                      "Aug", "Sep", "Oct", "Nov", "Dec"]


# ============================================================
# 1. FORMAT TANGGAL & ANGKA DWIBAHASA
# ============================================================

def _tgl_id(d: date) -> str:
    """'31 Juli 2026' -- format tanggal panjang Indonesia, sama persis
    gaya file referensi (nama bulan lengkap, tanpa singkatan)."""
    return f"{d.day} {_BULAN_ID[d.month - 1]} {d.year}"


def _tgl_en(d: date) -> str:
    """'July 31, 2026' -- gaya Amerika lengkap, dipakai di judul dokumen
    (baris 'July 31, And December 31, 2025')."""
    return f"{_BULAN_EN[d.month - 1]} {d.day}, {d.year}"


def _tgl_en_singkat(d: date) -> str:
    """'31 DEC 2025' -- dipakai sbg header kolom tabel (kapital semua,
    persis pola '31 DEC 2025' di file referensi)."""
    return f"{d.day} {_BULAN_EN_SINGKAT[d.month - 1].upper()} {d.year}"


def _tgl_id_singkat(d: date) -> str:
    """'31 JULI 2026' -- header kolom tabel sisi Indonesia (kapital)."""
    return f"{d.day} {_BULAN_ID[d.month - 1].upper()} {d.year}"


def _rp(v: Optional[float]) -> str:
    """Format Rupiah tanpa simbol, koma ribuan, negatif dalam kurung,
    nol/None jadi '-' -- PERSIS gaya file referensi (mis. kolom Deposito
    31 Dec 2025 yang kosong tampil '-', bukan '0' atau '0.00')."""
    if v is None:
        return "-"
    v = round(float(v), 0)
    if v == 0:
        return "-"
    if v < 0:
        return f"({abs(v):,.0f})"
    return f"{v:,.0f}"


# ============================================================
# 2. HELPER STYLING python-docx
# ============================================================

def _set_font(run, bold=False, italic=False, underline=False,
              size=_UKURAN_NORMAL, font_name=_FONT_UTAMA):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.color.rgb = _WARNA_HITAM
    # Font Asia/kompleks openpyxl-style perlu diset terpisah supaya Word
    # tidak diam-diam pakai font default lain utk sebagian karakter.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)


def _cell_text(cell, text: str, bold=False, italic=False, underline=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, size=_UKURAN_NORMAL,
               border_top=False, border_bottom=None):
    """Tulis teks ke 1 sel tabel dengan style seragam. border_bottom:
    None (default word)|'single'|'double' -- dipakai utk baris subtotal/
    total (garis di ATAS baris total, meniru gaya file referensi)."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    _set_font(run, bold=bold, italic=italic, underline=underline, size=size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if border_top:
        _atur_border_sel(cell, top="single")
    if border_bottom:
        _atur_border_sel(cell, bottom=border_bottom)


def _atur_border_sel(cell, top=None, bottom=None):
    """Set border atas/bawah 1 sel tabel (dipakai utk garis subtotal
    'Jumlah ...' -- python-docx tidak expose ini langsung, jadi lewat XML)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for sisi, gaya in (("top", top), ("bottom", bottom)):
        if gaya is None:
            continue
        el = OxmlElement(f'w:{sisi}')
        el.set(qn('w:val'), gaya)  # 'single' atau 'double'
        el.set(qn('w:sz'), "6" if gaya == "single" else "10")
        el.set(qn('w:space'), "1")
        el.set(qn('w:color'), "000000")
        borders.append(el)


def _hapus_semua_border_tabel(table):
    """Tabel data CALK tidak pakai garis kotak (beda dari tabel Excel) --
    hanya garis atas tipis di baris subtotal/total. python-docx bawaan
    tidak punya API 'no border' langsung, jadi override lewat XML."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for sisi in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f'w:{sisi}')
        el.set(qn('w:val'), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def _atur_lebar_kolom(table, lebar_cm: List[float]):
    """Set lebar kolom tabel secara eksplisit di setiap sel (python-docx
    kadang mengabaikan table.columns[i].width kalau tidak diulang per
    baris) -- juga matikan autofit supaya lebar manual ini benar2 dipakai."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(lebar_cm):
                cell.width = Cm(lebar_cm[idx])


def _tambah_field_halaman(paragraph, field_type: str):
    """Sisipkan field code Word (PAGE / NUMPAGES) -- dipakai utk footer
    'Halaman X dari Y' yang otomatis update kalau dibuka di Word."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' {field_type} '
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    _set_font(run, size=9)


# ============================================================
# 3. HEADER DOKUMEN & JUDUL NOTE (2 kolom dwibahasa)
# ============================================================

def _tambah_header_dokumen(doc: Document, nama_perusahaan: str,
                            tanggal_now: date, tanggal_lalu: date) -> None:
    """
    [FASE 4 -- diperbaiki] Blok kop, PERSIS pola file referensi:
        PT AKSHAYA ANANDA DYLAN LESTARI          (center, bold)
        CATATAN ATAS LAPORAN KEUANGAN | NOTES TO FINANCIAL STATEMENTS
        31 Juli 2026 Dan 31 Desember 2025 | July 31, And December 31, 2025
        (Disajikan dalam Rupiah...) | (Expressed in Rupiah...)
        ------------------------------------------------ (garis bawah)

    [PERBAIKAN Fase 4 poin 10 roadmap] Versi SEBELUMNYA menulis blok ini
    ke BODY dokumen (doc.add_paragraph()/doc.add_table()) -- itu cuma
    muncul 1x di halaman pertama, TIDAK berulang tiap halaman padahal
    file referensi (PT AADL, hal. 8/9/10/11) jelas menampilkan kop yang
    SAMA di SETIAP halaman. Versi ini menulis ke Word SECTION HEADER
    (doc.sections[0].header) supaya Word sendiri yang mengulanginya
    otomatis di semua halaman -- dipanggil SEKALI saja oleh
    susun_dan_tulis_semua_note_calk(), tidak perlu dipanggil ulang per
    note/per halaman.
    """
    header = doc.sections[0].header
    header.is_linked_to_previous = False

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nama_perusahaan.upper())
    _set_font(run, bold=True, size=12)
    p.paragraph_format.space_after = Pt(2)

    tabel = header.add_table(rows=3, cols=2, width=Cm(17))
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    _atur_lebar_kolom(tabel, [8.5, 8.5])

    baris = [
        ("CATATAN ATAS LAPORAN KEUANGAN", "NOTES TO FINANCIAL STATEMENTS", True, False),
        (f"{_tgl_id(tanggal_now)} Dan {_tgl_id(tanggal_lalu)}",
         f"{_tgl_en(tanggal_now)}, And {_tgl_en(tanggal_lalu)}", False, True),
        ("(Disajikan dalam Rupiah, kecuali dinyatakan lain)",
         "(Expressed in Rupiah, unless otherwise stated)", False, True),
    ]
    for r, (teks_id, teks_en, bold, italic) in enumerate(baris):
        _cell_text(tabel.cell(r, 0), teks_id, bold=bold, italic=italic,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_text(tabel.cell(r, 1), teks_en, bold=bold, italic=italic,
                   align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Garis pemisah tebal di bawah kop (paragraph border, BUKAN tabel --
    # lihat gotcha di docx SKILL.md: "jangan pakai tabel sbg horizontal
    # rule, pakai paragraph bottom border"). Ditambahkan ke HEADER juga
    # (bukan cuma body) supaya ikut berulang tiap halaman.
    p_garis = header.add_paragraph()
    p_garis.paragraph_format.space_before = Pt(2)
    p_garis.paragraph_format.space_after = Pt(6)
    _atur_border_paragraf(p_garis, bottom="single", sz=18)


def _atur_footer_nomor_halaman(doc: Document) -> None:
    """[FASE 4] Footer berulang tiap halaman: nomor halaman polos di
    tengah bawah -- PERSIS gaya file referensi (angka saja, mis. "8",
    "9", "10", "11" di kaki tiap halaman, TANPA label "Halaman X dari
    Y"). Pakai _tambah_field_halaman() yang sudah dibuat sebelumnya tapi
    belum pernah dipanggil di mana pun -- field code Word (PAGE) yang
    otomatis update sendiri, bukan angka statis yang bisa salah kalau
    urutan note berubah.
    """
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tambah_field_halaman(p, "PAGE")


def _atur_border_paragraf(paragraph, top=None, bottom=None, sz=6):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    for sisi, gaya in (("top", top), ("bottom", bottom)):
        if gaya is None:
            continue
        el = OxmlElement(f'w:{sisi}')
        el.set(qn('w:val'), gaya)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), "1")
        el.set(qn('w:color'), "000000")
        borders.append(el)
    p_pr.append(borders)


def _tambah_judul_note(doc: Document, nomor: str, judul_id: str, judul_en: str) -> None:
    """Baris judul note bernomor, gaya 2 kolom bold, mis.:
        3.   KAS DAN SETARA KAS   |   3.   CASH AND CASH EQUIVALENT
    Kolom Inggris SEDIKIT italic pada judulnya sendiri -- dicek di file
    referensi (mis. 'CASH AND CASH EQUIVALENT' bold, non-italic sebenarnya
    utk judul, italic dipakai di teks PENJELASAN di bawahnya saja -- lihat
    _tambah_paragraf_dwibahasa())."""
    tabel = doc.add_table(rows=1, cols=2)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    _atur_lebar_kolom(tabel, [8.5, 8.5])
    _cell_text(tabel.cell(0, 0), f"{nomor}.   {judul_id}", bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(tabel.cell(0, 1), f"{nomor}.   {judul_en}", bold=True,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)  # spacer tipis


def _tambah_subjudul_note(doc: Document, huruf: str, judul_id: str, judul_en: str) -> None:
    """Sub-item note (mis. 'e. Aset Tetap' / 'e. Fixed Assets') -- dipakai
    di Note 2 Ikhtisar Kebijakan Akuntansi yang punya banyak sub a-h."""
    tabel = doc.add_table(rows=1, cols=2)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    _atur_lebar_kolom(tabel, [8.5, 8.5])
    # [FIX -- bug titik dobel] `huruf` SELALU dikirim pemanggil dgn titik
    # sudah nempel (mis. "a.", "b.") -- lihat semua call site di file
    # ini. Sebelumnya baris ini nambahin ". " lagi di depan judul_id,
    # jadi tercetak "a.. Pendirian Entitas" (titik dobel) di SEMUA 8
    # sub-judul Note 1 & Note 2. Sisi EN di bawah TIDAK kena karena
    # tidak pernah nambah titik sendiri -- makanya cuma sisi ID yang
    # perlu diubah, bukan keduanya.
    _cell_text(tabel.cell(0, 0), f"{huruf} {judul_id}", bold=True, italic=True,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(tabel.cell(0, 1), f"{huruf} {judul_en}", bold=True, italic=True,
               align=WD_ALIGN_PARAGRAPH.LEFT)


def _tambah_paragraf_dwibahasa(doc: Document, teks_id: str, teks_en: str,
                                italic_en=True, spasi_setelah=6) -> None:
    """Paragraf penjelasan biasa (bukan tabel angka) -- kolom kiri
    Indonesia regular, kolom kanan Inggris italic (PERSIS gaya file
    referensi -- semua teks Inggris di badan note selalu italic, cuma
    judul note & label header tabel yang non-italic)."""
    tabel = doc.add_table(rows=1, cols=2)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    _atur_lebar_kolom(tabel, [8.5, 8.5])
    c_id, c_en = tabel.cell(0, 0), tabel.cell(0, 1)
    c_id.text = ""
    c_en.text = ""
    p_id = c_id.paragraphs[0]
    p_en = c_en.paragraphs[0]
    p_id.paragraph_format.space_after = Pt(spasi_setelah)
    p_en.paragraph_format.space_after = Pt(spasi_setelah)
    run_id = p_id.add_run(teks_id)
    _set_font(run_id, size=_UKURAN_NORMAL)
    run_en = p_en.add_run(teks_en)
    _set_font(run_en, italic=italic_en, size=_UKURAN_NORMAL)
    c_id.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    c_en.vertical_alignment = WD_ALIGN_VERTICAL.TOP


# ============================================================
# 4. TERJEMAHAN OTOMATIS (Anthropic API) -- HANYA utk teks BEBAS
# ============================================================
# [PENTING] Teks boilerplate standar (Note 1 & 2, label header tabel,
# label "Jumlah"/"Total" dst) SENGAJA ditulis manual dwibahasa langsung
# di kode (lihat _NOTE2_KEBIJAKAN di bawah), BUKAN lewat AI translate --
# ini teks hukum/akuntansi baku yang sama persis di SETIAP laporan CALK
# klien manapun, jadi terjemahan manual yang sudah direview sekali jauh
# lebih akurat & konsisten drpd manggil API tiap generate (juga lebih
# cepat & tidak kena biaya API berulang utk teks yang tidak pernah
# berubah). AI translate HANYA dipakai utk teks yang BENAR-BENAR bebas/
# spesifik per klien -- field "keterangan" di COA yang diisi manual
# akuntan (mis. catatan khusus di suatu akun) -- karena teks itu TIDAK
# BISA diprediksi/ditulis manual sebelumnya.

_CACHE_TERJEMAHAN_FILE = os.path.join(
    os.path.dirname(__file__), "..", "cache_terjemahan_calk.json"
)


def _muat_cache_terjemahan() -> Dict[str, str]:
    try:
        with open(_CACHE_TERJEMAHAN_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}


def _simpan_cache_terjemahan(cache: Dict[str, str]) -> None:
    try:
        with open(_CACHE_TERJEMAHAN_FILE, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
    except OSError as e:
        logger.warning(f"Gagal simpan cache terjemahan CALK: {e}")


_SYSTEM_PROMPT_TERJEMAHAN_CALK = (
    "Kamu penerjemah teknis akuntansi Indonesia->Inggris utk "
    "Catatan atas Laporan Keuangan (Notes to Financial "
    "Statements) perusahaan Indonesia. Terjemahkan teks "
    "berikut ke Inggris bisnis/akuntansi formal, gaya sama "
    "seperti laporan keuangan beraudit internasional. "
    "Balas HANYA dengan hasil terjemahan, tanpa embel2, "
    "tanpa tanda kutip, tanpa penjelasan tambahan."
)


def terjemahkan_id_ke_en(teks_id: str, cache: Optional[Dict[str, str]] = None) -> str:
    """
    Terjemahkan teks bebas Indonesia -> Inggris lewat Claude, KHUSUS utk
    teks yang tidak ada di boilerplate manual (mis. isi kolom
    "keterangan" custom per akun/klien). Dicache per-teks (persis
    string) supaya generate berikutnya dgn teks sama tidak manggil API
    lagi -- penting krn 1 laporan CALK bisa berisi puluhan baris
    keterangan, dan laporan yang sama sering di-generate ulang (mis.
    setelah 1 akun dikoreksi).

    Kalau ANTHROPIC_API_KEY tidak ada di environment / API gagal,
    FALLBACK ke teks asli (Indonesia) apa adanya + tag "[EN?]" di depan
    supaya kelihatan jelas di dokumen bahwa baris itu belum diterjemahkan
    (BUKAN diam-diam menampilkan teks Indonesia seolah itu versi Inggris,
    dan BUKAN membuat generate laporan gagal total hanya krn API key
    belum diisi/kuota habis -- laporan tetap harus bisa keluar).

    [FIX -- POINT 3/4 gaya] Sebelumnya fungsi ini menulis (flush)
    SELURUH file cache ke disk setelah SETIAP SATU terjemahan baru
    (_simpan_cache_terjemahan di akhir). Untuk 1 dokumen CALK dengan
    banyak baris "keterangan" custom, ini berarti puluhan write file
    sinkron + (lebih parah) puluhan HTTP request BERURUTAN/blocking ke
    Claude API selama generate 1 dokumen -- tiap request bisa 1-3 detik,
    jadi total delay bisa puluhan detik cuma utk terjemahan. Sekarang
    fungsi ini HANYA mengisi dict `cache` di memori, TIDAK menulis file
    apa pun -- caller (tulis_note_akun_generik/susun_dan_tulis_semua_
    note_calk) yang bertanggung jawab memanggil _simpan_cache_terjemahan()
    SEKALI SAJA di akhir, setelah semua note selesai diproses. Ini juga
    menghapus race condition read-modify-write file cache kalau 2
    generate CALK berjalan bersamaan (2 request server berbeda).

    [UBAH -- migrasi ke claude_client.py] SEBELUMNYA fungsi ini bikin
    panggilan HTTP mentah sendiri (requests.post ke api.anthropic.com,
    baca ANTHROPIC_API_KEY & model "claude-sonnet-5" hardcode terpisah)
    -- TERPISAH TOTAL dari modules/claude_client.py, titik masuk
    terpusat yang justru dibuat supaya SEMUA panggilan Claude di backend
    lewat SATU tempat (client caching, retry eksplisit utk 429/500/503/
    529, audit trail) -- lihat docstring di claude_client.py. Sekarang
    lewat claude_client.panggil_claude_teks() -- dapat retry otomatis &
    audit trail GRATIS tanpa kode tambahan di sini. Perilaku FALLBACK
    (API key belum diset / semua percobaan tetap gagal setelah retry)
    TETAP SAMA PERSIS seperti sebelumnya: kembalikan teks asli + tag
    "[EN?]" di depan, TIDAK menggagalkan generate CALK -- satu-satunya
    beda, "gagal" sekarang baru terjadi SETELAH retry (bukan 1x percobaan).
    """
    teks_id = (teks_id or "").strip()
    if not teks_id:
        return ""

    if cache is None:
        cache = _muat_cache_terjemahan()
    if teks_id in cache:
        return cache[teks_id]

    try:
        teks_en = panggil_claude_teks(
            teks_id,
            modul_pemanggil="calk_export",
            system_prompt=_SYSTEM_PROMPT_TERJEMAHAN_CALK,
            max_tokens=500,
        ).strip()
        if not teks_en:
            raise ValueError("Respons API kosong")
    except ClaudeError as e:
        # ClaudeError SUDAH mencakup kasus ANTHROPIC_API_KEY belum diset
        # (lihat claude_client.ambil_client()) -- pesan log dibedakan
        # sedikit supaya gampang dibedakan di log server, tapi hasil
        # fallback yang dikembalikan ke caller SAMA ("[EN?] ...").
        logger.warning(f"Auto-translate CALK gagal (Claude) utk '{teks_id[:40]}...': {e}")
        return f"[EN?] {teks_id}"
    except Exception as e:
        logger.warning(f"Auto-translate CALK gagal utk '{teks_id[:40]}...': {e}")
        return f"[EN?] {teks_id}"

    cache[teks_id] = teks_en
    return teks_en


# ============================================================
# 5. TABEL AKUN DWIBAHASA (inti Note 3 dst -- Kas/Piutang/dll)
# ============================================================
# Layout 4 kolom per baris, label menghadap tepi luar & angka menghadap
# tengah (PERSIS pola file referensi):
#   [label akun ID, rata kiri] [nilai periode SEKARANG, rata kanan] |
#   [nilai periode LALU, rata kanan] [label akun EN, rata kiri]

_LEBAR_KOLOM_TABEL_AKUN = [5.5, 3.0, 3.0, 5.5]  # cm, total 17 (~margin A4 wajar)


def _tabel_akun_baru(doc: Document) -> Any:
    tabel = doc.add_table(rows=0, cols=4)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    return tabel


def _tambah_baris_header_periode(tabel, label_now_id: str, label_now_en: str,
                                  label_lalu_id: str, label_lalu_en: str) -> None:
    """Baris '31 JULI 2026 | 31 DEC 2025' -- underline, bold, rata kanan
    (kolom label kiri/kanan dikosongkan, cuma 2 kolom tengah yang diisi)."""
    r = tabel.add_row()
    _atur_lebar_kolom_baris(r, _LEBAR_KOLOM_TABEL_AKUN)
    _cell_text(r.cells[0], "", align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(r.cells[1], label_now_id, bold=True, underline=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(r.cells[2], label_lalu_id, bold=True, underline=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(r.cells[3], "", align=WD_ALIGN_PARAGRAPH.LEFT)


def _atur_lebar_kolom_baris(row, lebar_cm: List[float]) -> None:
    for idx, cell in enumerate(row.cells):
        if idx < len(lebar_cm):
            cell.width = Cm(lebar_cm[idx])


def _tambah_baris_akun(tabel, label_id: str, nilai_now: Optional[float],
                        nilai_lalu: Optional[float], label_en: str,
                        indent=False, bold=False, garis_atas=False) -> None:
    """1 baris akun biasa, mis. 'Pety Cash - Putu | 6,381,181 | 110,298 |
    Pety Cash - Putu Soemantri'. `indent` -> label akun sedikit menjorok
    (dipakai utk item di bawah sub-grup, mis. akun bank di bawah grup
    'Omah Soemantri' pada Note Piutang -- lihat contoh gambar referensi
    Note 4). `garis_atas` -> garis tipis di atas baris ini (dipakai di
    baris 'Jumlah ...' persis sebelum subtotal)."""
    r = tabel.add_row()
    _atur_lebar_kolom_baris(r, _LEBAR_KOLOM_TABEL_AKUN)
    prefix = "    " if indent else ""
    _cell_text(r.cells[0], prefix + label_id, bold=bold,
               align=WD_ALIGN_PARAGRAPH.LEFT, border_top=garis_atas)
    _cell_text(r.cells[1], _rp(nilai_now), bold=bold,
               align=WD_ALIGN_PARAGRAPH.RIGHT, border_top=garis_atas)
    _cell_text(r.cells[2], _rp(nilai_lalu), bold=bold,
               align=WD_ALIGN_PARAGRAPH.RIGHT, border_top=garis_atas)
    _cell_text(r.cells[3], prefix + label_en, bold=bold,
               align=WD_ALIGN_PARAGRAPH.LEFT, border_top=garis_atas)


def _tambah_baris_subgrup(tabel, label_id: str, label_en: str) -> None:
    """Baris label sub-grup TANPA angka, mis. 'Omah Soemantri' sebelum
    daftar akun channel penjualannya (Booking.Com, Tiket.Com, dst) --
    lihat Note 4 Piutang Usaha di gambar referensi."""
    r = tabel.add_row()
    _atur_lebar_kolom_baris(r, _LEBAR_KOLOM_TABEL_AKUN)
    _cell_text(r.cells[0], label_id, align=WD_ALIGN_PARAGRAPH.LEFT)
    _cell_text(r.cells[1], "", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(r.cells[2], "", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(r.cells[3], label_en, align=WD_ALIGN_PARAGRAPH.LEFT)


def _tambah_baris_jumlah(tabel, label_id: str, nilai_now: float,
                          nilai_lalu: float, label_en: str,
                          garis_ganda=False) -> None:
    """Baris 'Jumlah ... | Total ...' -- bold, garis tipis di atas
    (garis_ganda=True utk baris 'Jumlah Kas dan Setara Kas' paling
    bawah/grand total note, meniru border-bottom double di file model
    Excel yang dipakai sheet lain -- konsisten satu keluarga gaya)."""
    r = tabel.add_row()
    _atur_lebar_kolom_baris(r, _LEBAR_KOLOM_TABEL_AKUN)
    gaya_border = "double" if garis_ganda else "single"
    _cell_text(r.cells[0], label_id, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
               border_top=True)
    _cell_text(r.cells[1], _rp(nilai_now), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
               border_top=True)
    _cell_text(r.cells[2], _rp(nilai_lalu), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
               border_top=True)
    _cell_text(r.cells[3], label_en, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
               border_top=True)
    if garis_ganda:
        for c in r.cells:
            _atur_border_sel(c, bottom="double")


def tulis_note_akun_generik(doc: Document, nomor: str, judul_id: str, judul_en: str,
                             daftar_akun: List[Dict[str, Any]],
                             tanggal_now: date, tanggal_lalu: date,
                             kalimat_pembuka_id: Optional[str] = None,
                             kalimat_pembuka_en: Optional[str] = None,
                             label_jumlah_id: Optional[str] = None,
                             label_jumlah_en: Optional[str] = None,
                             cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """
    Engine UTAMA yang dipakai berulang utk Note 3 (Kas), 4 (Piutang
    Usaha), 5 (Piutang Lainnya), dan SEMUA note kategori akun lain yang
    di-auto-generate oleh susun_dan_tulis_semua_note_calk() di bawah
    (Persediaan, Biaya Dibayar Dimuka, Aset Tetap, Utang, Modal,
    Pendapatan, Beban, dst) -- supaya 1 kali ditulis & diverifikasi
    benar, otomatis konsisten di SEMUA note angka, bukan copy-paste
    manual per note yang gampang divergen.

    Args:
        daftar_akun: list akun SUDAH dikelompokkan, tiap item salah satu:
            {"tipe": "akun", "label_id", "label_en", "now", "lalu",
             "indent": bool (opsional)}
            {"tipe": "subgrup", "label_id", "label_en"}          -- baris
                label tanpa angka (mis. "Omah Soemantri")
            {"tipe": "subtotal", "label_id", "label_en"}         -- [BARU]
                baris "Jumlah ..." DI TENGAH note (mis. "Jumlah Kas dan
                Bank" sebelum grup Deposito di Note 3 -- lihat file
                referensi). Nilainya = akumulasi item "akun" SEJAK
                subtotal sebelumnya (atau sejak awal note kalau ini
                subtotal pertama), dan counter subgrup-nya di-reset
                sesudahnya. Grand total di akhir note TETAP dihitung dari
                seluruh item "akun" (subtotal cuma checkpoint visual,
                tidak dobel-hitung).
        kalimat_pembuka_*: opsional, teks 1 baris sblm tabel (mis.
            "Rincian kas dan setara kas adalah sebagai berikut:") --
            kalau None, note langsung ke tabel tanpa kalimat pembuka
            (dipakai note yg tidak butuh, mis. note tanpa rincian akun).
        label_jumlah_*: [BARU] override label baris total PALING AKHIR.
            Default "Jumlah {judul_id}"/"Total {judul_en}" (title-case)
            -- tapi sebagian note di file referensi labelnya beda dari
            judul note-nya sendiri (mis. Note 4 Piutang Usaha labelnya
            cuma "Jumlah"/"Total" polos, bukan "Jumlah Piutang Usaha"),
            jadi disediakan override eksplisit per pemanggil.
        cache_terjemahan: [FIX -- POINT 3/4 gaya] opsional -- dict cache
            terjemahan SHARED, dikelola pemanggil (susun_dan_tulis_semua_
            note_calk, dipanggil 1x untuk SELURUH dokumen CALK, dioper ke
            tiap panggilan tulis_note_akun_generik yang ke-12 note).
            Sebelumnya tiap panggilan fungsi ini me-load cache dari file
            sendiri-sendiri (12x baca file per dokumen) dan menulis file
            ulang setelah tiap 1 terjemahan baru (puluhan write per
            dokumen). Kalau None (dipanggil berdiri sendiri/standalone,
            mis. dari test), fallback ke perilaku lama: load sendiri &
            simpan sendiri di akhir fungsi ini.
    """
    _tambah_judul_note(doc, nomor, judul_id, judul_en)
    if kalimat_pembuka_id:
        _tambah_paragraf_dwibahasa(doc, kalimat_pembuka_id,
                                    kalimat_pembuka_en or kalimat_pembuka_id,
                                    spasi_setelah=2)

    tabel = _tabel_akun_baru(doc)
    # [FIX] Panggilan lama salah kirim tanggal_now dua kali ke kolom
    # label_lalu_id, jadi header kolom "lalu" ikut nampilin "31 JULI
    # 2026" (harusnya "31 DEC 2025"). File referensi konsisten: kolom
    # "now" pakai format Indonesia (_tgl_id_singkat), kolom "lalu" pakai
    # format Inggris singkat (_tgl_en_singkat) -- lihat docstring kedua
    # fungsi itu. label_now_en/label_lalu_en tidak dipakai badan
    # _tambah_baris_header_periode() (cuma cells[1] & cells[2] yang
    # ditulis), jadi dikosongkan saja supaya tidak menyesatkan pembaca.
    _tambah_baris_header_periode(
        tabel, _tgl_id_singkat(tanggal_now), "",
        _tgl_en_singkat(tanggal_lalu), "",
    )

    # [FIX -- BUG NYATA #2] Sebelumnya terjemahkan_id_ke_en() adalah dead
    # code -- didokumentasikan di komentar atas (baris ~386) sebagai jalur
    # utk field "keterangan" per-akun di COA (AkunCoaRequest.keterangan,
    # lihat main.py), tapi tidak pernah benar-benar dipanggil di mana pun.
    # Sekarang diwire di sini: kalau item akun punya "keterangan" (teks
    # bebas yg diisi akuntan) TAPI belum ada "keterangan_en", auto-translate
    # lewat terjemahkan_id_ke_en() (cache + fallback "[EN?]" kalau API
    # gagal/key belum diset -- lihat docstring fungsi itu). Cache dimuat
    # SEKALI di sini (bukan per-item) supaya 1 note dgn banyak akun
    # berketerangan tidak buka file cache berulang-ulang.
    ada_keterangan = any(item.get("tipe", "akun") == "akun" and item.get("keterangan")
                          for item in daftar_akun)
    # [FIX -- POINT 3/4 gaya] Pakai cache yang dioper pemanggil kalau ada
    # (shared across 12 note dalam 1 dokumen CALK). Hanya load sendiri +
    # simpan sendiri di akhir kalau fungsi ini dipanggil standalone tanpa
    # cache dioper (mis. dipanggil langsung dari test/skrip lain) --
    # itulah gunanya flag `_cache_lokal` di bawah, supaya tidak
    # menyimpan cache milik pemanggil lain secara tidak sengaja.
    _cache_lokal = cache_terjemahan is None
    if cache_terjemahan is None:
        cache_terjemahan = _muat_cache_terjemahan() if ada_keterangan else None

    # [BARU -- PERBAIKAN PERFORMA -- PARALELISASI TERJEMAHAN] Sebelumnya
    # tiap "keterangan" akun yang butuh terjemahan dipanggil SATU PER
    # SATU secara berurutan (blocking) di dalam for-loop bawah -- kalau
    # 1 note punya banyak akun berketerangan, artinya banyak HTTP request
    # ke Claude API menunggu bergiliran (tiap request bisa 1-3 detik),
    # padahal tiap terjemahan itu independen satu sama lain (tidak saling
    # butuh hasil terjemahan lain). Sekarang teks-teks yang BELUM ada di
    # cache dikumpulkan dulu (dedup -- teks yang sama persis di beberapa
    # akun cukup diterjemahkan sekali), diterjemahkan PARALEL lewat
    # ThreadPoolExecutor (network I/O-bound, sama pola dengan
    # _panggil_ai_batch_json_claude di akuntansi_ai.py), lalu hasilnya
    # dimasukkan ke cache_terjemahan SEBELUM for-loop utama jalan -- jadi
    # for-loop bawah (yang urutannya harus tetap sekuensial krn menyusun
    # tabel doc secara berurutan) tinggal baca dari cache (instan, tanpa
    # HTTP request lagi) lewat terjemahan_id_ke_en() yang tidak berubah
    # sama sekali. Item yang SUDAH punya "keterangan_en" manual (tidak
    # butuh AI translate) otomatis tidak ikut dikumpulkan di sini.
    if ada_keterangan and cache_terjemahan is not None:
        teks_perlu_terjemahan = sorted({
            (item.get("keterangan") or "").strip()
            for item in daftar_akun
            if item.get("tipe", "akun") == "akun"
            and item.get("keterangan")
            and not item.get("keterangan_en")
            and (item.get("keterangan") or "").strip() not in cache_terjemahan
        })
        if teks_perlu_terjemahan:
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor(
                max_workers=min(6, len(teks_perlu_terjemahan))
            ) as executor:
                # cache={} (bukan None) supaya tiap thread TIDAK membaca
                # ulang file cache dari disk (sudah dimuat sekali di atas)
                # -- tetap konsisten dgn perbaikan "flush sekali di akhir"
                # yang didokumentasikan di docstring terjemahkan_id_ke_en().
                future_ke_teks = {
                    executor.submit(terjemahkan_id_ke_en, teks, {}): teks
                    for teks in teks_perlu_terjemahan
                }
                for future in concurrent.futures.as_completed(future_ke_teks):
                    teks = future_ke_teks[future]
                    try:
                        cache_terjemahan[teks] = future.result()
                    except Exception as e:  # noqa: BLE001 -- jangan gagalkan seluruh dokumen krn 1 terjemahan error
                        logger.warning(f"Auto-translate CALK (paralel) gagal utk '{teks[:40]}...': {e}")
                        cache_terjemahan[teks] = f"[EN?] {teks}"

    catatan_kaki_akun: List[Tuple[str, str, str]] = []  # (label_id, ket_id, ket_en)

    total_now = total_lalu = 0.0
    subtotal_now = subtotal_lalu = 0.0
    for item in daftar_akun:
        tipe = item.get("tipe", "akun")
        if tipe == "subgrup":
            _tambah_baris_subgrup(tabel, item["label_id"], item["label_en"])
        elif tipe == "subtotal":
            # Checkpoint visual di tengah note -- pakai akumulasi
            # subgrup (BUKAN total_now/total_lalu) supaya tidak
            # dobel-hitung ke grand total di akhir.
            _tambah_baris_jumlah(tabel, item["label_id"], subtotal_now,
                                  subtotal_lalu, item["label_en"])
            subtotal_now = subtotal_lalu = 0.0
        else:
            now = float(item.get("now") or 0.0)
            lalu = float(item.get("lalu") or 0.0)
            total_now += now
            total_lalu += lalu
            subtotal_now += now
            subtotal_lalu += lalu
            _tambah_baris_akun(tabel, item["label_id"], now, lalu,
                                item["label_en"], indent=item.get("indent", False))

            keterangan_id = (item.get("keterangan") or "").strip()
            if keterangan_id:
                keterangan_en = item.get("keterangan_en") or terjemahkan_id_ke_en(
                    keterangan_id, cache=cache_terjemahan,
                )
                catatan_kaki_akun.append((item["label_id"], keterangan_id, keterangan_en))

    label_j_id = label_jumlah_id or f"Jumlah {judul_id.title()}"
    label_j_en = label_jumlah_en or f"Total {judul_en.title()}"
    _tambah_baris_jumlah(tabel, label_j_id, total_now, total_lalu,
                          label_j_en, garis_ganda=True)

    # Catatan per-akun (kalau ada) ditulis SETELAH tabel angka -- format
    # tabel 4-kolom (label|now|lalu|label_en) tidak punya tempat wajar utk
    # teks bebas panjang, jadi ditulis sbg paragraf dwibahasa kecil,
    # 1 baris per akun yg punya keterangan, urut sesuai urutan tabel.
    for label_id, ket_id, ket_en in catatan_kaki_akun:
        _tambah_paragraf_dwibahasa(
            doc, f"Keterangan {label_id}: {ket_id}", f"Note {label_id}: {ket_en}",
            spasi_setelah=2,
        )

    # [FIX -- POINT 3/4 gaya] Flush cache ke disk HANYA kalau fungsi ini
    # dipanggil standalone (cache_terjemahan tidak dioper dari pemanggil).
    # Kalau dipanggil dari susun_dan_tulis_semua_note_calk() (mode
    # normal), flush dilakukan pemanggil SEKALI di akhir seluruh dokumen
    # -- lihat catatan di sana.
    if _cache_lokal and cache_terjemahan:
        _simpan_cache_terjemahan(cache_terjemahan)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)  # spacer antar-note


# ============================================================
# 6. NOTE 1 "UMUM" -- profil perusahaan (data BELUM ada di DB)
# ============================================================
# [PENTING] Client model di db_client.py SAAT INI cuma punya
# nama/lokasi/kontak -- field detail (akta pendirian, notaris, SK
# Kemenkumham, bidang usaha, susunan komisaris/direksi, jumlah karyawan)
# BELUM ada tabelnya. `profil` di bawah HARUS disuplai pemanggil dari
# sumber lain (mis. form isian khusus di frontend, atau tabel baru kalau
# nanti mau dipermanenkan) -- kalau field tertentu kosong, ditulis
# placeholder "-- lengkapi data --" supaya generate tetap jalan (tidak
# error), bukan diam-diam dikosongkan tanpa keterangan.

CONTOH_PROFIL: Dict[str, Any] = {
    "nama_perusahaan": "PT AKSHAYA ANANDA DYLAN LESTARI",
    "nomor_akta_pendirian": "15", "tanggal_akta_pendirian": "30 Desember 2022",
    "nama_notaris": "Yade Erianzah Waldo, S.H, M.Kn.",
    # [FIX] "Tahun 2022" sekarang bagian dari string ini sendiri (bukan
    # ditempel otomatis oleh tulis_note_1_umum() lagi -- lihat catatan
    # FIX di sana), konsisten dgn konvensi penomoran SK Kemenkumham asli.
    "no_sk_kemenkumham": "AHU-AH.01.09.0075428 Tahun 2022", "tanggal_sk": "10 November 2022",
    "nomor_akta_perubahan_terakhir": "1", "tanggal_akta_perubahan_terakhir": "2 November 2022",
    "no_sk_perubahan_terakhir": "AHU-0081552.AH.01.02.TAHUN 2022", "tanggal_sk_perubahan": "10 November 2022",
    "bidang_usaha_id": ("Rumah Minum/Kafe, Restoran, Real Estat Yang Dimiliki Sendiri Atau "
                          "Disewa, Hotel Bintang, Hotel Melati, Aktivitas Penyewaan dan Sewa "
                          "Guna Usaha Tanpa Hak Opsi Mobil, Bus, Truk, dan Sejenisnya, "
                          "Aktivitas Penyewaan dan Sewa Guna Usaha Tanpa Hak Opsi Alat Pesta, "
                          "Penyediaan Akomodasi Lainnya"),
    "bidang_usaha_en": ("Drinking Houses/Cafes, Restaurants, Owned or Rented Real Estate, "
                          "Star Hotels, Budget Hotels, Rental and Leasing Activities without "
                          "Option Rights for Cars, Buses, Trucks and the Like, Rental and "
                          "Leasing Activities without Option Rights for Party Equipment, "
                          "Provision of Other Accommodation"),
    "domisili_id": "Jakarta Selatan, Provinsi Jakarta", "domisili_en": "South Jakarta, Jakarta",
    "tahun_mulai_operasi": "2022",
    "komisaris": [("Komisaris Utama", "Ignatio Katriel Swy", "Main Commissioner"),
                  ("Komisaris", "Retno Wulan Sugihartati", "Commissioner")],
    "direksi": [("Direktur Utama", "Happy Februanty", "Main Director"),
                ("Direktur", "Bambang Widjonarko", "Director")],
    "kepala_cabang": [("Kepala Cabang Omah Soemantri", "Sanis Paweningsih",
                        "Omah Soemantri's Branch Head"),
                       ("Kepala Cabang Putu Soemantri", "Ida Bagus Putu Astawa S.",
                        "Putu Soemantri's Branch Head")],
    "jumlah_karyawan_lalu": 28, "tahun_karyawan_lalu": "2025",
    "jumlah_karyawan_now": 31, "tahun_karyawan_now": "2026",
    "umur_manfaat_inventaris": "4-8", "umur_manfaat_bangunan": "20",
}


def _isi(profil: Dict[str, Any], key: str, default="-- lengkapi data --") -> Any:
    v = (profil or {}).get(key)
    return v if v not in (None, "") else default


def tulis_note_1_umum(doc: Document, nomor: str, profil: Dict[str, Any],
                       tanggal_now: Optional[date] = None) -> None:
    """
    [FIX] tanggal_now ditambahkan sebagai parameter opsional supaya
    kalimat "Susunan anggota Dewan Komisaris dan Direksi Entitas pada
    tanggal ... adalah sebagai berikut:" bisa diisi tanggalnya --
    sebelumnya parameter ini tidak ada sama sekali sehingga kalimat
    selalu tercetak tanpa tanggal (lihat file referensi PT AADL: "pada
    tanggal 31 Juli 2026 adalah sebagai berikut"). Dibuat Optional
    (default None -> placeholder "-- lengkapi data --") supaya
    pemanggilan lama yang belum sempat diupdate tidak langsung error,
    tapi orchestrator (susun_dan_tulis_semua_note_calk) SUDAH diupdate
    mengirim tanggal_now, jadi placeholder ini seharusnya tidak pernah
    muncul lewat jalur normal.
    """
    p = _isi
    nama = p(profil, "nama_perusahaan")
    tgl_susunan_id = _tgl_id(tanggal_now) if tanggal_now else "-- lengkapi data --"
    tgl_susunan_en = _tgl_en(tanggal_now) if tanggal_now else "-- lengkapi data --"

    _tambah_judul_note(doc, nomor, "UMUM", "GENERAL")
    _tambah_subjudul_note(doc, "a.", "Pendirian Entitas", "Establishment Entity")

    teks_id_1 = (
        f'{nama} ("Perseroan") didirikan di Indonesia berdasarkan Akta '
        f'Notaris {p(profil,"nama_notaris")} No. {p(profil,"nomor_akta_pendirian")} '
        f'tanggal {p(profil,"tanggal_akta_pendirian")}. Akta Pendirian tersebut '
        f'telah mendapat pengesahan dari Kementerian Hukum dan Hak Asasi '
        f'Manusia Republik Indonesia dengan Surat Keputusan No. '
        # [FIX -- hardcode tahun] sebelumnya " Tahun 2022"/" Year 2022"
        # ditulis literal di sini (nyalin apa adanya dari dokumen
        # referensi PT AADL, bukan diparameterkan) -- salah untuk SEMUA
        # client lain & tahun lain, dan dobel kalau no_sk_kemenkumham
        # yg diisi user SUDAH memuat "Tahun 20XX" sendiri (konvensi
        # penomoran SK Kemenkumham Indonesia memang begitu -- lihat
        # CONTOH_PROFIL & hint field di CalkProfilForm.jsx). Sekarang
        # no_sk_kemenkumham dipakai APA ADANYA, user yang mengisi
        # nomor SK lengkap dgn "Tahun 20XX" kalau memang ada di
        # dokumen aslinya.
        f'{p(profil,"no_sk_kemenkumham")} tanggal {p(profil,"tanggal_sk")}.'
    )
    teks_en_1 = (
        f'{nama} ("the Company") was established in Indonesia based on a '
        f'Deed of Notary {p(profil,"nama_notaris")} No. {p(profil,"nomor_akta_pendirian")} '
        f'dated {p(profil,"tanggal_akta_pendirian")}. The Deed of Establishment has '
        f'been approved by the Ministry of Law and Human Rights of the '
        f'Republic of Indonesia with Decree No. {p(profil,"no_sk_kemenkumham")} '
        f'dated {p(profil,"tanggal_sk")}.'
    )
    _tambah_paragraf_dwibahasa(doc, teks_id_1, teks_en_1)

    teks_id_2 = (
        f'Anggaran dasar Perseroan untuk pertama kali mengalami perubahan, '
        f'berdasarkan Akta Notaris {p(profil,"nama_notaris")} No. '
        f'{p(profil,"nomor_akta_perubahan_terakhir")} tanggal '
        f'{p(profil,"tanggal_akta_perubahan_terakhir")} tentang Pernyataan '
        f'Keputusan Para Pemegang Saham Yang Diedarkan {nama} Sebagai '
        f'Pengganti Rapat Umum Pemegang Saham Luar Biasa. Perubahan ini '
        f'telah diterima oleh Kementerian Hukum dan Hak Asasi Manusia '
        f'Republik Indonesia dalam Surat Keputusan No. '
        f'{p(profil,"no_sk_perubahan_terakhir")} tanggal {p(profil,"tanggal_sk_perubahan")}.'
    )
    teks_en_2 = (
        f'The Company\'s articles of association underwent changes for the '
        f'first time, based on the Deed of Notary {p(profil,"nama_notaris")} No. '
        f'{p(profil,"nomor_akta_perubahan_terakhir")} dated '
        f'{p(profil,"tanggal_akta_perubahan_terakhir")} concerning Statement of '
        f'Shareholders\' Decisions Circulated by {nama} in lieu of the '
        f'Extraordinary General Meeting of Shareholders. This change has '
        f'been accepted by the Ministry of Law and Human Rights of the '
        f'Republic of Indonesia in Decree No. {p(profil,"no_sk_perubahan_terakhir")} '
        f'dated {p(profil,"tanggal_sk_perubahan")}.'
    )
    _tambah_paragraf_dwibahasa(doc, teks_id_2, teks_en_2)

    teks_id_3 = (
        f'Sesuai dengan Pasal 2 Anggaran Dasar Perusahaan, Perseroan '
        f'bergerak di bidang {p(profil,"bidang_usaha_id")}.'
    )
    teks_en_3 = (
        f'In accordance with Article 2 of the Company\'s Articles of '
        f'Association, the Company is engaged in the field of '
        f'{p(profil,"bidang_usaha_en")}.'
    )
    _tambah_paragraf_dwibahasa(doc, teks_id_3, teks_en_3)

    teks_id_4 = (
        f'Perusahaan berdomisili di {p(profil,"domisili_id")} dan mulai '
        f'beroperasi secara komersial pada tahun {p(profil,"tahun_mulai_operasi")}.'
    )
    teks_en_4 = (
        f'The Company is domiciled in {p(profil,"domisili_en")} and started '
        f'its commercial operations in {p(profil,"tahun_mulai_operasi")}.'
    )
    _tambah_paragraf_dwibahasa(doc, teks_id_4, teks_en_4)

    _tambah_subjudul_note(doc, "b.", "Susunan Komisaris dan Direksi",
                           "The Board of Commissioner and Directors")
    _tambah_paragraf_dwibahasa(
        doc,
        f"Susunan anggota Dewan Komisaris dan Direksi Entitas pada tanggal "
        f"{tgl_susunan_id} adalah sebagai berikut:",
        f"The members of the Board of Commissioner and Directors Entity as "
        f"of {tgl_susunan_en} are as follows:",
        spasi_setelah=2,
    )

    tabel = doc.add_table(rows=0, cols=4)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    _atur_lebar_kolom(tabel, [4.0, 4.5, 4.5, 4.0])

    def _baris_jabatan(jabatan_id, nama_orang, jabatan_en, header=False):
        r = tabel.add_row()
        _atur_lebar_kolom_baris(r, [4.0, 4.5, 4.5, 4.0])
        _cell_text(r.cells[0], jabatan_id, bold=header, underline=header)
        _cell_text(r.cells[1], nama_orang, bold=header, underline=header)
        _cell_text(r.cells[2], "", bold=header)
        _cell_text(r.cells[3], jabatan_en, bold=header, underline=header,
                   italic=not header, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _baris_jabatan("Komisaris", "", "Commissioner", header=True)
    for jabatan_id, nama_orang, jabatan_en in profil.get("komisaris", []):
        _baris_jabatan(jabatan_id, nama_orang, jabatan_en)
    _baris_jabatan("Direksi", "", "Director", header=True)
    for jabatan_id, nama_orang, jabatan_en in profil.get("direksi", []):
        _baris_jabatan(jabatan_id, nama_orang, jabatan_en)
    for jabatan_id, nama_orang, jabatan_en in profil.get("kepala_cabang", []):
        _baris_jabatan(jabatan_id, nama_orang, jabatan_en)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    teks_id_kry = (
        f'Pada tahun {p(profil,"tahun_karyawan_lalu")} Perseroan memiliki '
        f'{p(profil,"jumlah_karyawan_lalu")} karyawan, dan pada tahun '
        f'{p(profil,"tahun_karyawan_now")} memiliki {p(profil,"jumlah_karyawan_now")} karyawan.'
    )
    teks_en_kry = (
        f'As of {p(profil,"tahun_karyawan_lalu")} the Company has '
        f'{p(profil,"jumlah_karyawan_lalu")} employees, and as of '
        f'{p(profil,"tahun_karyawan_now")} the Company has {p(profil,"jumlah_karyawan_now")} Employees.'
    )
    _tambah_paragraf_dwibahasa(doc, teks_id_kry, teks_en_kry)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ============================================================
# 7. NOTE 2 "IKHTISAR KEBIJAKAN AKUNTANSI" -- boilerplate SAK ETAP
# ============================================================
# Teks a-d & f di bawah SAMA PERSIS utk semua klien SAK ETAP (baku
# standar akuntansi, bukan spesifik 1 perusahaan) -- cuma sub "e. Aset
# Tetap" yang punya angka umur manfaat spesifik per klien (diambil dari
# `profil`, konsisten dgn kebijakan penyusutan aset tetap klien ybs).

def tulis_note_2_kebijakan_akuntansi(doc: Document, nomor: str,
                                      profil: Dict[str, Any]) -> None:
    _tambah_judul_note(doc, nomor, "IKHTISAR KEBIJAKAN AKUNTANSI",
                        "SUMMARY OF SIGNIFICANT ACCOUNTING POLICIES")

    _tambah_subjudul_note(doc, "a.", "Dasar Penyusunan Laporan Keuangan",
                           "Basis of Preparation of the Financial Statements")
    _tambah_paragraf_dwibahasa(
        doc,
        "Laporan Keuangan Perusahaan telah disajikan sesuai dengan Standar "
        "Akuntansi Keuangan untuk Entitas Tanpa Akuntabilitas Publik (SAK "
        "ETAP) di Indonesia dan disusun berdasarkan prinsip berkesinambungan "
        "(going concern) serta mengikuti konversi harga historis (historical "
        "cost). Mata uang pelaporan yang digunakan untuk penyusunan laporan "
        "keuangan adalah mata uang rupiah (Rp.) Kebijakan akuntansi ini "
        "diterapkan secara konsisten kecuali apabila dinyatakan adanya "
        "perubahan dalam kebijakan akuntansi yang dianut.",
        "The Company's Financial Statements have been prepared in "
        "accordance with the Financial Accounting Standards for Entities "
        "Without Public Accountability (SAK ETAP) in Indonesia and have "
        "been prepared based on the principle of going concern and "
        "following historical cost conversions. The reporting currency "
        "used for the preparation of the financial statements is the "
        "Rupiah currency (Rp.). This accounting policy is applied "
        "consistently unless there is a change in the adopted accounting "
        "policy.",
    )
    _tambah_paragraf_dwibahasa(
        doc,
        "Laporan arus kas menyajikan penerimaan dan pembayaran kas dan bank "
        "yang diklasifikasikan ke dalam aktivitas operasi, investasi dan "
        "pendanaan, dengan arus kas dari aktivitas operasi disajikan dengan "
        "menggunakan metode langsung.",
        "The statement of cash flows presents the receipts and payments of "
        "cash and banks classified into operating, investing and financing "
        "activities, with cash flows from operating activities presented "
        "using the direct method.",
    )
    _tambah_paragraf_dwibahasa(
        doc,
        "Periode pelaporan keuangan Perusahaan adalah 1 Januari - 31 "
        "Desember. Mata uang pelaporan yang digunakan dalam laporan "
        "keuangan adalah Rupiah, yang merupakan mata uang fungsional "
        "Perusahaan.",
        "The financial reporting period of the Company is January 1 - "
        "December 31. The reporting currency used in the financial "
        "statements is Rupiah, which is the Company's functional currency.",
    )

    _tambah_subjudul_note(doc, "b.", "Kas dan Setara Kas", "Cash and Cash Equivalent")
    _tambah_paragraf_dwibahasa(
        doc,
        "Kas dan setara kas terdiri dari kas, bank, dan deposito berjangka "
        "yang tidak dijaminkan dan tidak dibatasi penggunaannya.",
        "Cash and cash equivalents consist of cash on hand and in banks "
        "which are not pledged as collateral and are not restricted.",
    )

    _tambah_subjudul_note(doc, "c.", "Piutang Usaha", "Account Receivable")
    _tambah_paragraf_dwibahasa(
        doc,
        "Perusahaan tidak melakukan penyisihan atas kemungkinan tidak "
        "tertagihnya piutang usaha.",
        "The Company does not provide allowance for the possibility of "
        "uncollectible account receivable.",
    )

    _tambah_subjudul_note(doc, "d.", "Biaya Dibayar Dimuka", "Prepaid Expenses")
    _tambah_paragraf_dwibahasa(
        doc,
        "Biaya dibayar di muka diamortisasi dan dibebankan pada usaha "
        "selama masa manfaatnya dengan menggunakan metode garis lurus dan "
        "disajikan sebagai aset lancar atau aset tidak lancar berdasarkan "
        "sifatnya.",
        "Prepaid expenses are amortized and charged to operations over "
        "their beneficial periods using the straight-line method and "
        "presented as current asset or non-current asset based on their "
        "nature.",
    )

    # --- Halaman baru (biasanya) -- sub e & f dilanjutkan; python-docx
    # tidak WAJIB page-break manual di sini krn Word akan wrap otomatis,
    # tapi disediakan util add_page_break() kalau pemanggil mau kontrol
    # penuh 1 note per halaman (lihat susun_calk_docx()).
    _tambah_subjudul_note(doc, "e.", "Aset Tetap", "Fixed Assets")
    _tambah_paragraf_dwibahasa(
        doc,
        "Aset tetap dinyatakan sebesar biaya perolehan dikurangi akumulasi "
        "penyusutan dan rugi penurunan nilai. Biaya perolehan termasuk "
        "biaya penggantian bagian aset tetap pada saat biaya tersebut "
        "terjadi, jika kriteria pengakuan terpenuhi. Demikian juga, ketika "
        "inspeksi besar dilakukan, biayanya diakui dalam jumlah tercatat "
        "aset sebagai penggantian jika memenuhi kriteria pengakuan. Semua "
        "biaya perbaikan dan pemeliharaan lainnya yang tidak memenuhi "
        "kriteria pengakuan diakui dalam laporan laba rugi dan penghasilan "
        "komprehensif lain pada saat terjadinya.",
        "Fixed assets are stated at cost less accumulated depreciation and "
        "impairment losses. Cost includes the cost of replacing part of "
        "the fixed assets when that cost is incurred, if the recognition "
        "criteria are met. Likewise, when a major inspection is performed, "
        "its cost is recognized in the carrying amount of the fixed assets "
        "as a replacement if the recognition criteria are satisfied. All "
        "other repairs and maintenance costs that do not meet the "
        "recognition criteria are recognized in statement of profit or "
        "loss and other comprehensive income as incurred.",
    )
    _tambah_paragraf_dwibahasa(
        doc,
        "Penyusutan aset dimulai pada saat tersedia untuk digunakan dan "
        "dihitung dengan menggunakan metode garis lurus berdasarkan "
        "taksiran masa manfaat aset sebagai berikut:",
        "Depreciation of an asset starts when it is available for use and "
        "is computed using the straight-line method based on the "
        "estimated useful lives of the assets as follows:",
        spasi_setelah=2,
    )

    tabel = doc.add_table(rows=0, cols=4)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    _atur_lebar_kolom(tabel, [5.0, 3.5, 3.5, 5.0])
    r = tabel.add_row()
    _atur_lebar_kolom_baris(r, [5.0, 3.5, 3.5, 5.0])
    _cell_text(r.cells[0], "", underline=True)
    _cell_text(r.cells[1], "Tahun", underline=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(r.cells[2], "", underline=True)
    _cell_text(r.cells[3], "Years", underline=True, italic=True,
               align=WD_ALIGN_PARAGRAPH.RIGHT)

    # [FIX -- hardcode client referensi] Sebelumnya label EN nempel
    # f"Inventaris - {profil.get('nama_cabang_utama','Omah Soemantri')}"
    # -- field "nama_cabang_utama" TIDAK PERNAH ada di CalkProfilRequest
    # (main.py) atau CalkProfilForm.jsx, jadi fallback "Omah Soemantri"
    # (nama cabang PT AADL, klien referensi) SELALU yang kepakai, untuk
    # SEMUA client lain juga. Sekarang label ID & EN dibuat simetris,
    # generik (tanpa nama cabang), sampai ada kebutuhan nyata utk
    # field itu -- kalau nanti perlu, tambah field resmi ke
    # CalkProfilRequest dulu, jangan hardcode nama client lagi di sini.
    daftar_umur_manfaat = [
        ("Inventaris", profil.get("umur_manfaat_inventaris", "4-8"), "Inventaris"),
        ("Bangunan", profil.get("umur_manfaat_bangunan", "20"), "Building"),
    ]
    for label_id, tahun, label_en in daftar_umur_manfaat:
        r = tabel.add_row()
        _atur_lebar_kolom_baris(r, [5.0, 3.5, 3.5, 5.0])
        _cell_text(r.cells[0], label_id)
        _cell_text(r.cells[1], str(tahun), align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(r.cells[2], "")
        _cell_text(r.cells[3], label_en, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    _tambah_paragraf_dwibahasa(
        doc,
        "Suatu aset tetap dihentikan pengakuannya pada saat dilepaskan "
        "atau ketika tidak ada lagi manfaat ekonomis masa depan yang "
        "diharapkan dari penggunaan atau pelepasannya. Keuntungan atau "
        "kerugian yang timbul dari penghentian pengakuan aset (dihitung "
        "sebagai selisih antara hasil neto hasil pelepasan dan nilai "
        "tercatat aset) dimasukkan dalam laporan laba rugi dan penghasilan "
        "komprehensif lain pada tahun aset tersebut dihentikan pengakuannya.",
        "An item of fixed assets is derecognized upon disposal or when no "
        "future economic benefits are expected from its use or disposal. "
        "Any gain or loss arising on derecognition of the asset "
        "(calculated as the difference between the net disposal proceeds "
        "and the carrying amount of the asset) is included in statement "
        "of profit or loss and other comprehensive income in the year the "
        "asset is derecognized.",
    )
    _tambah_paragraf_dwibahasa(
        doc,
        "Nilai residu aset tetap, masa manfaat dan metode penyusutan "
        "ditelaah, dan disesuaikan secara prospektif jika sesuai, pada "
        "setiap akhir tahun buku.",
        "The fixed asset's residual values, useful lives and methods of "
        "depreciation are reviewed, and adjusted prospectively if "
        "appropriate, at each financial year end.",
    )

    _tambah_subjudul_note(doc, "f.", "Pengakuan Pendapatan dan Beban",
                           "Revenue and Expense Recognition")
    _tambah_paragraf_dwibahasa(
        doc,
        "Pendapatan diakui pada saat jasa telah direalisasikan, baik pada "
        "yang sudah maupun yang belum dibuatkan invoice. Sedangkan biaya "
        "dibebankan pada periode yang sama dengan pendapatan yang "
        "bersangkutan. Sesuai dengan asas \"matching of cost against the "
        "revenue\", Beban diakui pada saat terjadinya dan sesuai dengan "
        "masa manfaatnya (accrual basis).",
        "Revenue is accrued when services have been realized, both for "
        "invoices that have been made and have not been made. While costs "
        "are charged in the same period as the revenue. In accordance "
        "with the principle of \"matching of costs against the revenue\", "
        "expenses are recognized when incurred and according to their "
        "useful life (accrual basis).",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ============================================================
# 7. NOTE 3 "KAS DAN SETARA KAS" -- 2 subtotal (Kas & Bank, Deposito)
# ============================================================

def tulis_note_3_kas_dan_setara_kas(doc: Document, nomor: str,
                                     daftar_kas_bank: List[Dict[str, Any]],
                                     daftar_deposito: List[Dict[str, Any]],
                                     tanggal_now: date, tanggal_lalu: date,
                                     cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note 3 Kas dan Setara Kas -- PERSIS pola file referensi: baris
    petty cash + rekening bank, ditutup subtotal 'Jumlah Kas dan Bank',
    lanjut baris deposito berjangka, ditutup subtotal 'Jumlah Deposito',
    baru grand total 'Jumlah Kas dan Setara Kas' paling bawah (garis
    dobel).

    Args:
        daftar_kas_bank: list item kas kecil & rekening bank, tiap item
            {"label_id", "label_en", "now", "lalu"} (TANPA "tipe" --
            fungsi ini yang nempelin tipe "akun" + subtotal-nya sendiri)
        daftar_deposito: sama, utk deposito berjangka. Kalau client
            tidak punya deposito, KIRIM LIST KOSONG ([]) -- bukan
            di-skip -- supaya baris "Jumlah Deposito" tetap muncul dgn
            nilai "-", konsisten dgn file referensi (31 DEC 2025 kolom
            Deposito tetap tampil "-" walau nol).
    """
    daftar_akun = (
        [{"tipe": "akun", **item} for item in daftar_kas_bank]
        + [{"tipe": "subtotal", "label_id": "Jumlah Kas dan Bank",
            "label_en": "Total Cash and Bank"}]
        + [{"tipe": "akun", **item} for item in daftar_deposito]
        + [{"tipe": "subtotal", "label_id": "Jumlah Deposito",
            "label_en": "Total Deposito"}]
    )
    tulis_note_akun_generik(
        doc, nomor, "KAS DAN SETARA KAS", "CASH AND CASH EQUIVALENT",
        daftar_akun, tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian kas dan setara kas adalah sebagai berikut:",
        kalimat_pembuka_en="The details of cash and cash equivalent are as follows:",
        label_jumlah_id="Jumlah Kas dan Setara Kas",
        label_jumlah_en="Total Cash and Cash Equivalent",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 8. NOTE 4 "PIUTANG USAHA" -- campuran subgrup+indent & item datar
# ============================================================

def tulis_note_4_piutang_usaha(doc: Document, nomor: str,
                                daftar_piutang: List[Dict[str, Any]],
                                tanggal_now: date, tanggal_lalu: date,
                                cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note 4 Piutang Usaha -- PERSIS pola file referensi: sebagian
    cabang dipecah per channel penjualan (subgrup label "Omah Soemantri"
    lalu item indented "Booking.Com"/"Tiket.Com"/dst di bawahnya),
    sebagian cabang langsung 1 baris nilai (mis. "Warung Soemantri"
    tanpa breakdown channel). Label total di file referensi cuma
    "Jumlah"/"Total" polos (BUKAN "Jumlah Piutang Usaha"), makanya
    di-override.

    Args:
        daftar_piutang: list item, tiap item salah satu:
            {"tipe": "subgrup", "label_id", "label_en"} -- header cabang
                yg mau dipecah channel (mis. "Omah Soemantri")
            {"tipe": "akun", "label_id", "label_en", "now", "lalu",
             "indent": True}  -- item channel di bawah subgrup di atas
            {"tipe": "akun", "label_id", "label_en", "now", "lalu"}
                (indent=False/tidak diisi) -- cabang tanpa breakdown
                channel, mis. "Warung Soemantri", "Café Kapu Soemantri"
        Urutan list HARUS sama dgn urutan tampil yg diinginkan (subgrup
        diikuti item indent-nya, baru cabang datar berikutnya).
    """
    tulis_note_akun_generik(
        doc, nomor, "PIUTANG USAHA", "ACCOUNT RECEIVABLES",
        daftar_piutang, tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian piutang usaha adalah sebagai berikut:",
        kalimat_pembuka_en="The details of trade receivables are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 9. NOTE 5 "PIUTANG LAINNYA" -- biasanya 1 baris polos
# ============================================================

def tulis_note_5_piutang_lainnya(doc: Document, nomor: str,
                                  daftar_piutang_lain: List[Dict[str, Any]],
                                  tanggal_now: date, tanggal_lalu: date,
                                  cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note 5 Piutang Lainnya -- di file referensi cuma 1 baris "Piutang
    Lain-Lain" tanpa rincian lebih jauh, tapi dibuat generik (terima
    list) kalau suatu saat client punya lebih dari 1 jenis piutang lain
    (mis. "Piutang Karyawan", "Piutang Pemegang Saham", dst -- tinggal
    tambah item ke list, tidak perlu ubah fungsi ini).

    Args:
        daftar_piutang_lain: list item {"label_id", "label_en", "now",
            "lalu"} (TANPA "tipe" -- fungsi ini yang nempelin "akun").
    """
    daftar_akun = [{"tipe": "akun", **item} for item in daftar_piutang_lain]
    tulis_note_akun_generik(
        doc, nomor, "PIUTANG LAINNYA", "OTHER RECEIVABLE",
        daftar_akun, tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian piutang lain-lain adalah sebagai berikut:",
        kalimat_pembuka_en="The details of other receivables are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 10. NOTE 6 "BIAYA DIBAYAR DIMUKA"
# ============================================================

def tulis_note_6_biaya_dibayar_dimuka(doc: Document, nomor: str,
                                       daftar_biaya: List[Dict[str, Any]],
                                       tanggal_now: date, tanggal_lalu: date,
                                       cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Biaya Dibayar Dimuka -- pola akun datar standar (sewa
    dibayar dimuka, asuransi dibayar dimuka, dll), sama seperti Note 5.

    Args:
        daftar_biaya: list item {"label_id","label_en","now","lalu"}.
    """
    daftar_akun = [{"tipe": "akun", **item} for item in daftar_biaya]
    tulis_note_akun_generik(
        doc, nomor, "BIAYA DIBAYAR DIMUKA", "PREPAID EXPENSES",
        daftar_akun, tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian biaya dibayar dimuka adalah sebagai berikut:",
        kalimat_pembuka_en="The details of prepaid expenses are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 11. NOTE 7 "ASET TETAP" -- tabel mutasi (Saldo Awal/Penambahan/
#     Pengurangan/Saldo Akhir), BEDA struktur dari note akun biasa
# ============================================================
# [PENTING] Tidak ada contoh gambar dari user utk note ini -- struktur
# di bawah mengikuti FORMAT STANDAR SAK ETAP utk tabel mutasi aset tetap
# (dipakai hampir semua CALK di Indonesia: kolom Saldo Awal/Penambahan/
# Pengurangan/Saldo Akhir, dipisah 2 blok "Biaya Perolehan" & "Akumulasi
# Penyusutan", ditutup baris "Nilai Buku"). Gaya visual (font, label ID
# kiri/EN kanan, angka ditengah) tetap DIPERTAHANKAN konsisten dgn note
# lain, cuma jumlah kolom angkanya 4 (bukan 2) krn butuh mutasi, bukan
# cuma 2 titik waktu.

_LEBAR_KOLOM_TABEL_ASET = [3.4, 2.35, 2.35, 2.35, 2.35, 3.4]  # cm, total 16.2


def _tabel_aset_tetap_baru(doc: Document) -> Any:
    tabel = doc.add_table(rows=0, cols=6)
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
    _hapus_semua_border_tabel(tabel)
    return tabel


def _tambah_header_mutasi_aset(tabel) -> None:
    """Baris header 'Saldo Awal | Penambahan | Pengurangan | Saldo Akhir'
    (ID di atas, EN italic di baris kedua DALAM sel yang sama, supaya
    tetap 1 baris tabel tapi 2 bahasa -- beda dari note akun biasa yang
    pisah kolom kiri/kanan, karena di sini kolom label akun sudah dipakai
    utk nama kategori aset, bukan label header)."""
    r = tabel.add_row()
    _atur_lebar_kolom_baris(r, _LEBAR_KOLOM_TABEL_ASET)
    _cell_text(r.cells[0], "", underline=True)
    for idx, (lbl_id, lbl_en) in enumerate([
        ("Saldo Awal", "Beginning Balance"),
        ("Penambahan", "Addition"),
        ("Pengurangan", "Deduction"),
        ("Saldo Akhir", "Ending Balance"),
    ], start=1):
        cell = r.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run1 = p.add_run(lbl_id)
        _set_font(run1, bold=True, underline=True, size=9)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(lbl_en)
        _set_font(run2, bold=True, italic=True, underline=True, size=9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_text(r.cells[5], "", underline=True)


def _tambah_baris_mutasi_aset(tabel, label_id: str, label_en: str,
                               saldo_awal: float, penambahan: float,
                               pengurangan: float, saldo_akhir: float,
                               bold=False, garis_atas=False) -> None:
    r = tabel.add_row()
    _atur_lebar_kolom_baris(r, _LEBAR_KOLOM_TABEL_ASET)
    _cell_text(r.cells[0], label_id, bold=bold, align=WD_ALIGN_PARAGRAPH.LEFT,
               size=9.5, border_top=garis_atas)
    for idx, nilai in enumerate([saldo_awal, penambahan, pengurangan, saldo_akhir], start=1):
        _cell_text(r.cells[idx], _rp(nilai), bold=bold,
                   align=WD_ALIGN_PARAGRAPH.RIGHT, size=9.5, border_top=garis_atas)
    _cell_text(r.cells[5], label_en, bold=bold, align=WD_ALIGN_PARAGRAPH.LEFT,
               size=9.5, border_top=garis_atas)


def tulis_note_7_aset_tetap(doc: Document, nomor: str,
                             daftar_kategori: List[Dict[str, Any]],
                             nilai_buku_now: float, nilai_buku_lalu: float,
                             tanggal_now: date, tanggal_lalu: date,
                             beban_penyusutan_tahun_berjalan: Optional[float] = None) -> None:
    """Note Aset Tetap -- 2 blok tabel mutasi (Biaya Perolehan &
    Akumulasi Penyusutan) per kategori aset, ditutup baris Nilai Buku.

    Args:
        daftar_kategori: list per kategori aset (Bangunan, Inventaris,
            Kendaraan, dll), tiap item:
            {"label_id", "label_en",
             "biaya_saldo_awal", "biaya_penambahan", "biaya_pengurangan", "biaya_saldo_akhir",
             "akum_saldo_awal", "akum_penambahan", "akum_pengurangan", "akum_saldo_akhir"}
            (penambahan akumulasi penyusutan = beban penyusutan tahun
            berjalan kategori itu; pengurangan = akumulasi penyusutan
            aset yang dilepas/dijual)
        nilai_buku_now/lalu: total nilai buku (biaya - akumulasi) periode
            berjalan & pembanding -- dihitung pemanggil dari data aset,
            BUKAN dihitung ulang di sini (supaya 1 sumber kebenaran sama
            dgn angka Nilai Buku yang tampil di Neraca).
        beban_penyusutan_tahun_berjalan: opsional, kalau diisi ditambah
            1 kalimat narasi jumlah beban penyusutan tahun ini (standar
            di CALK aset tetap).
    """
    _tambah_judul_note(doc, nomor, "ASET TETAP", "FIXED ASSETS")
    _tambah_paragraf_dwibahasa(
        doc,
        "Rincian aset tetap dan akumulasi penyusutan adalah sebagai berikut:",
        "The details of fixed assets and accumulated depreciation are as follows:",
        spasi_setelah=2,
    )

    # --- Blok 1: Biaya Perolehan / Cost ---
    p = doc.add_paragraph()
    run = p.add_run("Biaya Perolehan")
    _set_font(run, bold=True, italic=True, size=_UKURAN_NORMAL)
    run2 = p.add_run("  /  Cost")
    _set_font(run2, bold=True, italic=True, size=_UKURAN_NORMAL)

    tabel_biaya = _tabel_aset_tetap_baru(doc)
    _tambah_header_mutasi_aset(tabel_biaya)
    tot_biaya = [0.0, 0.0, 0.0, 0.0]
    for kat in daftar_kategori:
        vals = [kat.get("biaya_saldo_awal", 0) or 0, kat.get("biaya_penambahan", 0) or 0,
                kat.get("biaya_pengurangan", 0) or 0, kat.get("biaya_saldo_akhir", 0) or 0]
        for i, v in enumerate(vals):
            tot_biaya[i] += float(v)
        _tambah_baris_mutasi_aset(tabel_biaya, kat["label_id"], kat["label_en"], *vals)
    _tambah_baris_mutasi_aset(tabel_biaya, "Jumlah Biaya Perolehan",
                               "Total Cost", *tot_biaya, bold=True, garis_atas=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # --- Blok 2: Akumulasi Penyusutan / Accumulated Depreciation ---
    p = doc.add_paragraph()
    run = p.add_run("Akumulasi Penyusutan")
    _set_font(run, bold=True, italic=True, size=_UKURAN_NORMAL)
    run2 = p.add_run("  /  Accumulated Depreciation")
    _set_font(run2, bold=True, italic=True, size=_UKURAN_NORMAL)

    tabel_akum = _tabel_aset_tetap_baru(doc)
    _tambah_header_mutasi_aset(tabel_akum)
    tot_akum = [0.0, 0.0, 0.0, 0.0]
    for kat in daftar_kategori:
        vals = [kat.get("akum_saldo_awal", 0) or 0, kat.get("akum_penambahan", 0) or 0,
                kat.get("akum_pengurangan", 0) or 0, kat.get("akum_saldo_akhir", 0) or 0]
        for i, v in enumerate(vals):
            tot_akum[i] += float(v)
        _tambah_baris_mutasi_aset(tabel_akum, kat["label_id"], kat["label_en"], *vals)
    _tambah_baris_mutasi_aset(tabel_akum, "Jumlah Akumulasi Penyusutan",
                               "Total Accumulated Depreciation", *tot_akum,
                               bold=True, garis_atas=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # --- Nilai Buku / Net Book Value (2 kolom periode, gaya note biasa) ---
    tabel_nb = _tabel_akun_baru(doc)
    _tambah_baris_header_periode(
        tabel_nb, _tgl_id_singkat(tanggal_now), "", _tgl_en_singkat(tanggal_lalu), "",
    )
    _tambah_baris_jumlah(tabel_nb, "Nilai Buku", nilai_buku_now, nilai_buku_lalu,
                          "Net Book Value", garis_ganda=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if beban_penyusutan_tahun_berjalan is not None:
        _tambah_paragraf_dwibahasa(
            doc,
            f"Beban penyusutan aset tetap tahun berjalan adalah sebesar "
            f"Rp {_rp(beban_penyusutan_tahun_berjalan)} yang dibebankan pada "
            f"laporan laba rugi dan penghasilan komprehensif lain.",
            f"Depreciation expense for the current year amounted to "
            f"Rp {_rp(beban_penyusutan_tahun_berjalan)} which was charged to "
            f"the statement of profit or loss and other comprehensive income.",
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ============================================================
# 12. NOTE 8 "UTANG USAHA"
# ============================================================

def tulis_note_8_utang_usaha(doc: Document, nomor: str,
                              daftar_utang: List[Dict[str, Any]],
                              tanggal_now: date, tanggal_lalu: date,
                              cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Utang Usaha -- pola sama dgn Note 4 Piutang Usaha (boleh
    campur subgrup+indent kalau utang dipecah per supplier/cabang, atau
    akun datar polos)."""
    tulis_note_akun_generik(
        doc, nomor, "UTANG USAHA", "TRADE PAYABLES",
        daftar_utang, tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian utang usaha adalah sebagai berikut:",
        kalimat_pembuka_en="The details of trade payables are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 13. NOTE 9 "UTANG PAJAK"
# ============================================================

def tulis_note_9_utang_pajak(doc: Document, nomor: str,
                              daftar_pajak: List[Dict[str, Any]],
                              tanggal_now: date, tanggal_lalu: date,
                              cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Utang Pajak -- rincian per jenis pajak (PPh 21, PPh 23, PPN
    Keluaran, PPh Final UMKM, dll -- SESUAI TaxCategory yang relevan
    dgn client, lihat modules/schemas.py). Ini note yang plg mudah
    disambung otomatis krn aplikasi ini SUDAH punya modul
    pph_badan.py/tax_router.py yang tahu jenis & jumlah pajak terutang
    client per periode.

    Args:
        daftar_pajak: list item {"label_id","label_en","now","lalu"},
            label_id biasanya nama jenis pajak (mis. "PPh Pasal 21",
            "PPh Final UMKM (PP 55/2022)", "PPN Keluaran").
    """
    tulis_note_akun_generik(
        doc, nomor, "UTANG PAJAK", "TAX PAYABLES",
        [{"tipe": "akun", **item} for item in daftar_pajak],
        tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian utang pajak adalah sebagai berikut:",
        kalimat_pembuka_en="The details of tax payables are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 14. NOTE 10 "BIAYA YANG MASIH HARUS DIBAYAR"
# ============================================================

def tulis_note_10_biaya_masih_harus_dibayar(doc: Document, nomor: str,
                                             daftar_biaya: List[Dict[str, Any]],
                                             tanggal_now: date, tanggal_lalu: date,
                                             cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Biaya yang Masih Harus Dibayar (accrued expenses) -- mis.
    gaji karyawan belum dibayar, listrik/air belum ditagih, dll."""
    tulis_note_akun_generik(
        doc, nomor, "BIAYA YANG MASIH HARUS DIBAYAR", "ACCRUED EXPENSES",
        [{"tipe": "akun", **item} for item in daftar_biaya],
        tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian biaya yang masih harus dibayar adalah sebagai berikut:",
        kalimat_pembuka_en="The details of accrued expenses are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 15. NOTE 11 "MODAL"
# ============================================================

def tulis_note_11_modal(doc: Document, nomor: str,
                         daftar_pemegang_saham: List[Dict[str, Any]],
                         tanggal_now: date, tanggal_lalu: date,
                         kalimat_tambahan_id: Optional[str] = None,
                         kalimat_tambahan_en: Optional[str] = None,
                         cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Modal -- rincian kepemilikan saham per pemegang saham (nilai
    setoran modal, BUKAN persentase -- kalau mau tampilkan lembar saham/
    persentase juga, tambahkan sbg kalimat_tambahan atau perluas fungsi
    ini nanti sesuai kebutuhan akta client).

    Args:
        daftar_pemegang_saham: list item {"label_id","label_en","now","lalu"}
        kalimat_tambahan_*: opsional, narasi tambahan setelah tabel (mis.
            "Tidak ada perubahan struktur permodalan selama periode berjalan.")
    """
    tulis_note_akun_generik(
        doc, nomor, "MODAL", "SHARE CAPITAL",
        [{"tipe": "akun", **item} for item in daftar_pemegang_saham],
        tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian modal disetor adalah sebagai berikut:",
        kalimat_pembuka_en="The details of paid-up capital are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )
    if kalimat_tambahan_id:
        _tambah_paragraf_dwibahasa(doc, kalimat_tambahan_id,
                                    kalimat_tambahan_en or kalimat_tambahan_id)


# ============================================================
# 16. NOTE 12 "PENDAPATAN USAHA"
# ============================================================

def tulis_note_12_pendapatan_usaha(doc: Document, nomor: str,
                                    daftar_pendapatan: List[Dict[str, Any]],
                                    tanggal_now: date, tanggal_lalu: date,
                                    cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Pendapatan Usaha -- [PENTING beda dari note Neraca lain]
    kolom periode di sini BUKAN 'posisi per tanggal neraca' tapi
    'akumulasi periode berjalan' (mis. '1 Jan s.d 31 Juli 2026' vs
    '1 Jan s.d 31 Des 2025' kalau laporan interim, atau 2 tahun penuh
    kalau laporan tahunan) -- karena Pendapatan/Beban itu akun laba
    rugi (flow), bukan akun neraca (stock). Header tanggal tetap dikirim
    lewat tanggal_now/tanggal_lalu spt note lain (dipakai apa adanya di
    _tgl_id_singkat/_tgl_en_singkat), TAPI pemanggil (orchestrator)
    WAJIB kirim nilai now/lalu yang sudah dihitung sbg akumulasi
    periode berjalan (dari susun_laba_rugi()), bukan saldo per tanggal.
    Boleh dipecah per channel/cabang penjualan sama seperti Note 4
    Piutang (pakai tipe 'subgrup' di list kalau perlu).
    """
    tulis_note_akun_generik(
        doc, nomor, "PENDAPATAN USAHA", "OPERATING REVENUE",
        daftar_pendapatan, tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian pendapatan usaha adalah sebagai berikut:",
        kalimat_pembuka_en="The details of operating revenue are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 17. NOTE 13 "BEBAN USAHA"
# ============================================================

def tulis_note_13_beban_usaha(doc: Document, nomor: str,
                               daftar_beban: List[Dict[str, Any]],
                               tanggal_now: date, tanggal_lalu: date,
                               cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Beban Usaha -- akun laba rugi (flow, sama catatan penting spt
    Note 12 di atas), biasanya dipecah per kategori beban (Gaji &
    Tunjangan, Sewa, Listrik/Air/Internet, Penyusutan, Perlengkapan,
    dll) -- list akun datar, urut sesuai besar/urutan COA."""
    tulis_note_akun_generik(
        doc, nomor, "BEBAN USAHA", "OPERATING EXPENSES",
        [{"tipe": "akun", **item} for item in daftar_beban],
        tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian beban usaha adalah sebagai berikut:",
        kalimat_pembuka_en="The details of operating expenses are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 18. NOTE 14 "PENDAPATAN (BEBAN) LAIN-LAIN"
# ============================================================

def tulis_note_14_pendapatan_beban_lain(doc: Document, nomor: str,
                                         daftar_lain: List[Dict[str, Any]],
                                         tanggal_now: date, tanggal_lalu: date,
                                         cache_terjemahan: Optional[Dict[str, str]] = None) -> None:
    """Note Pendapatan (Beban) Lain-lain -- akun di luar operasional
    utama (pendapatan jasa giro/bunga bank, beban administrasi bank,
    laba/rugi selisih kurs, laba/rugi pelepasan aset, dll). Nilai
    pendapatan biasanya positif, beban NEGATIF (supaya baris "Jumlah"
    otomatis jadi neto pendapatan (beban) lain-lain -- ikuti konvensi
    tanda yang sama dgn susun_laba_rugi())."""
    tulis_note_akun_generik(
        doc, nomor, "PENDAPATAN (BEBAN) LAIN-LAIN", "OTHER INCOME (EXPENSES)",
        [{"tipe": "akun", **item} for item in daftar_lain],
        tanggal_now, tanggal_lalu,
        kalimat_pembuka_id="Rincian pendapatan (beban) lain-lain adalah sebagai berikut:",
        kalimat_pembuka_en="The details of other income (expenses) are as follows:",
        label_jumlah_id="Jumlah", label_jumlah_en="Total",
        cache_terjemahan=cache_terjemahan,
    )


# ============================================================
# 19. NOTE 15 "PERPAJAKAN" -- rekonsiliasi fiskal & PPh Badan terutang
# ============================================================
# [PENTING] Beda dari note lain: fungsi ini TIDAK minta parameter mentah
# (label/now/lalu manual) tapi langsung terima dict `hasil_pph_badan`
# PERSIS output modules.pph_badan.hitung_pph_pasal_31e() -- supaya
# orchestrator tinggal oper hasil yg SUDAH dihitung tax_router/main.py
# (endpoint pph-badan/generate), tanpa transformasi ulang & tanpa risiko
# angka CALK menyimpang dari SPT Tahunan yg sesungguhnya dilaporkan.
#
# Kalau client pakai skema PPh Final UMKM (PP 55/2022, BUKAN Pasal
# 17/31E), field2 di atas (fasilitas_31e, dst) tidak relevan -- pakai
# tulis_note_15_perpajakan_final_umkm() di bawah sbg gantinya.

def tulis_note_15_perpajakan(doc: Document, nomor: str,
                              hasil_pph_badan: Dict[str, Any],
                              tanggal_now: date, tanggal_lalu: date,
                              pph_badan_terutang_lalu: Optional[float] = None) -> None:
    """Note Perpajakan skema Tarif Umum Pasal 17/31E -- rekonsiliasi
    laba komersial ke fiskal, lalu perhitungan PPh Badan terutang.

    Args:
        hasil_pph_badan: dict, PERSIS return value
            modules.pph_badan.hitung_pph_pasal_31e() periode berjalan.
        pph_badan_terutang_lalu: opsional, angka PPh Badan terutang
            periode pembanding (kalau ada/relevan ditampilkan) -- kalau
            None, kolom itu ditulis "-" (belum dihitung / tidak
            diaudit ulang, BUKAN otomatis dianggap nol).
    """
    rek = hasil_pph_badan.get("rekonsiliasi_fiskal", {})
    _tambah_judul_note(doc, nomor, "PERPAJAKAN", "TAXATION")
    _tambah_paragraf_dwibahasa(
        doc,
        "Rekonsiliasi antara laba komersial dan laba fiskal serta "
        "perhitungan Pajak Penghasilan Badan terutang adalah sebagai berikut:",
        "The reconciliation between commercial and fiscal income and the "
        "calculation of Corporate Income Tax payable are as follows:",
        spasi_setelah=2,
    )

    tabel = _tabel_akun_baru(doc)
    _tambah_baris_header_periode(
        tabel, _tgl_id_singkat(tanggal_now), "", _tgl_en_singkat(tanggal_lalu), "",
    )
    baris_rekon = [
        ("Laba Sebelum Pajak Menurut Laporan Laba Rugi Komersial",
         "Profit Before Tax per Commercial Statement of Profit or Loss",
         rek.get("laba_bersih_komersial", 0)),
        ("Koreksi Fiskal Positif", "Positive Fiscal Correction",
         rek.get("koreksi_fiskal_positif", 0)),
        ("Koreksi Fiskal Negatif", "Negative Fiscal Correction",
         -abs(rek.get("koreksi_fiskal_negatif", 0) or 0)),
    ]
    for label_id, label_en, nilai in baris_rekon:
        _tambah_baris_akun(tabel, label_id, nilai, None, label_en)
    _tambah_baris_jumlah(tabel, "Penghasilan Neto Fiskal", rek.get("penghasilan_neto_fiskal", 0),
                          None, "Net Fiscal Income")
    if rek.get("kompensasi_kerugian_fiskal"):
        _tambah_baris_akun(tabel, "Kompensasi Kerugian Fiskal",
                            -abs(rek["kompensasi_kerugian_fiskal"]), None,
                            "Fiscal Loss Compensation")
    _tambah_baris_jumlah(tabel, "Penghasilan Kena Pajak", rek.get("penghasilan_kena_pajak", 0),
                          None, "Taxable Income", garis_ganda=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # [FIX] modules.pph_badan.hitung_pph_pasal_31e() menyimpan
    # "persentase_pengurangan_pasal_31e" sbg PECAHAN (0.5 = 50%, lihat
    # PERSENTASE_PENGURANGAN_PASAL_31E di pph_badan.py), BUKAN sbg angka
    # persen siap pakai -- sebelumnya kode ini format langsung dgn
    # ":.0f}%" tanpa dikali 100, jadi tercetak "0%" (harusnya "50%").
    # Default fallback juga disamakan jadi 0.5 (pecahan) supaya
    # konsisten dgn satuan yang benar kalau field ini kosong.
    persentase_31e = (hasil_pph_badan.get('persentase_pengurangan_pasal_31e') or 0.5) * 100
    _tambah_paragraf_dwibahasa(
        doc,
        f"Perusahaan mendapat fasilitas pengurangan tarif sebesar "
        f"{persentase_31e:.0f}% "
        f"sesuai Pasal 31E Undang-Undang Pajak Penghasilan atas bagian "
        f"Penghasilan Kena Pajak yang memperoleh fasilitas.",
        f"The Company obtained a tax rate reduction facility of "
        f"{persentase_31e:.0f}% "
        f"in accordance with Article 31E of the Income Tax Law on the "
        f"portion of Taxable Income entitled to the facility.",
        spasi_setelah=2,
    )

    tabel2 = _tabel_akun_baru(doc)
    _tambah_baris_akun(tabel2, "Pajak Penghasilan Badan Terutang",
                        hasil_pph_badan.get("pph_badan_terutang", 0),
                        pph_badan_terutang_lalu, "Corporate Income Tax Payable",
                        bold=True)
    kredit = hasil_pph_badan.get("kredit_pajak", {}) or {}
    if kredit.get("total"):
        _tambah_baris_akun(tabel2, "Kredit Pajak", -abs(kredit["total"]), None,
                            "Tax Credit")
    label_status_id = ("Pajak Penghasilan Badan Kurang Bayar (Pasal 29)"
                        if hasil_pph_badan.get("status") == "KURANG BAYAR"
                        else "Pajak Penghasilan Badan Lebih Bayar (Pasal 28A)")
    label_status_en = ("Corporate Income Tax Underpayment (Article 29)"
                        if hasil_pph_badan.get("status") == "KURANG BAYAR"
                        else "Corporate Income Tax Overpayment (Article 28A)")
    nilai_status = (hasil_pph_badan.get("pph_pasal_29_kurang_bayar")
                     or -abs(hasil_pph_badan.get("pph_pasal_28a_lebih_bayar", 0) or 0))
    _tambah_baris_jumlah(tabel2, label_status_id, nilai_status, None,
                          label_status_en, garis_ganda=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _tambah_paragraf_dwibahasa(
        doc,
        "Perhitungan Pajak Penghasilan Badan di atas merupakan estimasi "
        "manajemen dan akan disesuaikan dengan Surat Pemberitahuan (SPT) "
        "Tahunan Pajak Penghasilan Badan yang dilaporkan.",
        "The above Corporate Income Tax calculation represents "
        "management's estimate and will be adjusted to the Corporate "
        "Annual Income Tax Return (SPT) as reported.",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tulis_note_15_perpajakan_final_umkm(doc: Document, nomor: str,
                                         peredaran_bruto_now: float,
                                         pph_final_now: float,
                                         peredaran_bruto_lalu: Optional[float],
                                         pph_final_lalu: Optional[float],
                                         tanggal_now: date, tanggal_lalu: date,
                                         tarif: float = 0.005) -> None:
    """Note Perpajakan skema PPh FINAL UMKM (PP 55 Tahun 2022 jo. PP 23
    Tahun 2018) -- dipakai kalau client BUKAN skema Tarif Umum Pasal
    17/31E (lihat catatan skema_pajak di pph_badan.py). Jauh lebih
    sederhana: PPh Final = tarif x peredaran bruto, TIDAK ADA
    rekonsiliasi fiskal krn PPh Final tidak mengenal koreksi fiskal.

    Args:
        tarif: default 0.005 (0,5%, tarif umum PP 55/2022). Perhatikan
            batas waktu penggunaan tarif final ini (badan usaha PT
            terbatas 3 tahun pajak sejak terdaftar) -- validasi periode
            berlaku dilakukan pemanggil, BUKAN di fungsi ini.
    """
    _tambah_judul_note(doc, nomor, "PERPAJAKAN", "TAXATION")
    _tambah_paragraf_dwibahasa(
        doc,
        f"Perusahaan dikenakan Pajak Penghasilan bersifat final sebesar "
        f"{tarif*100:.1f}% dari peredaran bruto sesuai Peraturan Pemerintah "
        f"Nomor 55 Tahun 2022 tentang Penyesuaian Pengaturan di Bidang "
        f"Pajak Penghasilan.",
        f"The Company is subject to final Income Tax of {tarif*100:.1f}% of "
        f"gross turnover in accordance with Government Regulation Number 55 "
        f"of 2022 concerning Adjustment of Income Tax Regulations.",
        spasi_setelah=2,
    )
    tabel = _tabel_akun_baru(doc)
    _tambah_baris_header_periode(
        tabel, _tgl_id_singkat(tanggal_now), "", _tgl_en_singkat(tanggal_lalu), "",
    )
    _tambah_baris_akun(tabel, "Peredaran Bruto", peredaran_bruto_now,
                        peredaran_bruto_lalu, "Gross Turnover")
    _tambah_baris_jumlah(
        tabel, f"Pajak Penghasilan Final ({tarif*100:.1f}%)", pph_final_now,
        pph_final_lalu, f"Final Income Tax ({tarif*100:.1f}%)", garis_ganda=True,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ============================================================
# 20. NOTE PENUTUP -- narasi saja (tanpa tabel akun): Pihak Berelasi,
#     Peristiwa Setelah Tanggal Neraca, Persetujuan Laporan Keuangan
# ============================================================

def tulis_note_pihak_berelasi(doc: Document, nomor: str,
                               daftar_transaksi_id: Optional[List[str]] = None,
                               daftar_transaksi_en: Optional[List[str]] = None) -> None:
    """Note Transaksi dengan Pihak-Pihak Berelasi -- kalau tidak ada
    transaksi berelasi material, tulis kalimat standar "tidak ada".
    Kalau ADA (mis. pinjaman dari pemegang saham, sewa dari pihak
    terafiliasi), isi daftar_transaksi_id/en (list kalimat, 1:1 index
    dgn versi Inggrisnya -- BUKAN diterjemahkan otomatis di sini,
    supaya penjelasan transaksi berelasi selalu direview manusia dulu
    krn sensitif buat SPT/audit)."""
    _tambah_judul_note(doc, nomor, "TRANSAKSI DENGAN PIHAK-PIHAK BERELASI",
                        "TRANSACTIONS WITH RELATED PARTIES")
    if daftar_transaksi_id:
        for teks_id, teks_en in zip(daftar_transaksi_id,
                                     daftar_transaksi_en or daftar_transaksi_id):
            _tambah_paragraf_dwibahasa(doc, teks_id, teks_en)
    else:
        _tambah_paragraf_dwibahasa(
            doc,
            "Dalam kegiatan usaha normal, Perusahaan tidak melakukan "
            "transaksi dengan pihak-pihak yang mempunyai hubungan istimewa "
            "yang bersifat material dan wajib diungkapkan.",
            "In the normal course of business, the Company did not enter "
            "into material transactions with related parties that are "
            "required to be disclosed.",
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tulis_note_peristiwa_setelah_tanggal_neraca(
        doc: Document, nomor: str, tanggal_neraca: date,
        daftar_peristiwa_id: Optional[List[str]] = None,
        daftar_peristiwa_en: Optional[List[str]] = None) -> None:
    """Note Peristiwa Setelah Tanggal Neraca (subsequent events) -- kalau
    tidak ada peristiwa material, tulis kalimat standar "tidak ada"."""
    _tambah_judul_note(doc, nomor, "PERISTIWA SETELAH TANGGAL NERACA",
                        "SUBSEQUENT EVENTS")
    if daftar_peristiwa_id:
        for teks_id, teks_en in zip(daftar_peristiwa_id,
                                     daftar_peristiwa_en or daftar_peristiwa_id):
            _tambah_paragraf_dwibahasa(doc, teks_id, teks_en)
    else:
        _tambah_paragraf_dwibahasa(
            doc,
            f"Tidak terdapat peristiwa material yang terjadi setelah "
            f"tanggal neraca ({_tgl_id(tanggal_neraca)}) sampai dengan "
            f"tanggal laporan keuangan ini diselesaikan, yang memerlukan "
            f"penyesuaian atau pengungkapan dalam laporan keuangan ini.",
            f"There were no material events occurring after the balance "
            f"sheet date ({_tgl_en(tanggal_neraca)}) through the date these "
            f"financial statements were completed, that require adjustment "
            f"or disclosure in these financial statements.",
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def tulis_note_persetujuan_laporan_keuangan(
        doc: Document, nomor: str, tanggal_persetujuan: date,
        nama_penanggung_jawab_id: str = "Direksi",
        nama_penanggung_jawab_en: str = "Board of Directors") -> None:
    """Note Persetujuan dan Penerbitan Laporan Keuangan -- SELALU note
    PALING TERAKHIR di CALK (standar SAK ETAP/SAK EMKM), menyatakan
    kapan & siapa yang bertanggung jawab menyetujui penerbitan laporan.
    [FIX] nama_penanggung_jawab dipecah jadi versi ID/EN terpisah --
    versi lama pakai 1 string yg sama utk 2 bahasa ("Direksi" nongol
    apa adanya di kolom Inggris, harusnya "Board of Directors")."""
    _tambah_judul_note(doc, nomor, "PERSETUJUAN DAN PENERBITAN LAPORAN KEUANGAN",
                        "APPROVAL AND ISSUANCE OF THE FINANCIAL STATEMENTS")
    _tambah_paragraf_dwibahasa(
        doc,
        f"Laporan keuangan ini telah disetujui dan diotorisasi untuk "
        f"terbit oleh {nama_penanggung_jawab_id} Perusahaan pada tanggal "
        f"{_tgl_id(tanggal_persetujuan)}.",
        f"These financial statements were approved and authorized for "
        f"issuance by the Company's {nama_penanggung_jawab_en} on "
        f"{_tgl_en(tanggal_persetujuan)}.",
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# === MULAI TEMPEL DI SINI ===

# ============================================================
# 21. ORCHESTRATOR [FASE 3 -- roadmap CALK] -- susun & tulis SEMUA
#     note CALK berurutan, auto-numbering, TIDAK menghitung ulang
#     saldo apa pun (100% pass-through dari Fase 1 & Fase 2)
# ============================================================

_FUNGSI_NOTE_STANDAR: Dict[str, Any] = {
    # note dgn signature standar (doc, nomor, daftar_akun, tanggal_now,
    # tanggal_lalu) -- SEMUA fungsi tulis_note_* di sini menerima list
    # item yg SUDAH punya key "tipe" (dari kelompokkan_akun_calk()),
    # jadi tinggal dioper langsung tanpa transformasi tambahan.
    "tulis_note_4_piutang_usaha": tulis_note_4_piutang_usaha,
    "tulis_note_5_piutang_lainnya": tulis_note_5_piutang_lainnya,
    "tulis_note_6_biaya_dibayar_dimuka": tulis_note_6_biaya_dibayar_dimuka,
    "tulis_note_8_utang_usaha": tulis_note_8_utang_usaha,
    "tulis_note_9_utang_pajak": tulis_note_9_utang_pajak,
    "tulis_note_10_biaya_masih_harus_dibayar": tulis_note_10_biaya_masih_harus_dibayar,
    "tulis_note_11_modal": tulis_note_11_modal,
    "tulis_note_12_pendapatan_usaha": tulis_note_12_pendapatan_usaha,
    "tulis_note_13_beban_usaha": tulis_note_13_beban_usaha,
    "tulis_note_14_pendapatan_beban_lain": tulis_note_14_pendapatan_beban_lain,
}


def _pisah_kas_bank_dan_deposito(
        daftar_akun: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """Note 3 (Kas dan Setara Kas) butuh 2 list terpisah utk 2 subtotal
    ("Jumlah Kas dan Bank" vs "Jumlah Deposito" -- lihat
    tulis_note_3_kas_dan_setara_kas()), tapi
    calk_mapping.kelompokkan_akun_calk() cuma balikin 1 list gabungan
    utk key "kas" (semua akun sub_kategori "Kas" di COA, tidak
    dibedakan). Dipisah di sini lewat KEYWORD nama akun ("deposito" di
    nama_akun -> Deposito, selainnya -> Kas & Bank) -- pola SAMA PERSIS
    dgn Keputusan #1 di calk_mapping.py (Piutang Usaha vs Lainnya).

    [CAVEAT -- sama seperti Keputusan #1] kalau ada akun deposito yang
    penamaannya di COA TIDAK mengandung kata "deposito" (mis. cuma
    "BCA Berjangka 5785"), item itu akan salah masuk ke Kas & Bank.
    Aman (akun tetap tampil, cuma di grup yang mungkin kurang tepat),
    tapi tetap wajib direview akuntan kalau ada kejanggalan subtotal.
    """
    kas_bank, deposito = [], []
    for item in daftar_akun:
        target = deposito if "deposito" in item["label_id"].lower() else kas_bank
        target.append({k: v for k, v in item.items() if k != "tipe"})
    return kas_bank, deposito


def susun_dan_tulis_semua_note_calk(
    doc: Document,
    profil: Dict[str, Any],
    neraca_now: Dict[str, Any], neraca_lalu: Dict[str, Any],
    laba_rugi_now: Dict[str, Any], laba_rugi_lalu: Dict[str, Any],
    tanggal_now: date, tanggal_lalu: date,
    aset_tetap: Optional[Dict[str, Any]] = None,
    hasil_pph_badan: Optional[Dict[str, Any]] = None,
    pph_badan_terutang_lalu: Optional[float] = None,
    pph_final_umkm: Optional[Dict[str, Any]] = None,
    pihak_berelasi: Optional[Dict[str, List[str]]] = None,
    peristiwa_setelah_neraca: Optional[Dict[str, List[str]]] = None,
    tanggal_persetujuan: Optional[date] = None,
    nama_penanggung_jawab_id: str = "Direksi",
    nama_penanggung_jawab_en: str = "Board of Directors",
    grouping_piutang_usaha: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    [FASE 3] Orchestrator utama: tulis header dokumen + SEMUA note CALK
    ke `doc` (Document() BARU/kosong yang disiapkan pemanggil), berurutan
    & auto-numbering mengikuti note mana yang benar-benar ada isinya.
    TIDAK doc.save() atau convert PDF -- itu tanggung jawab Fase 4
    (export_calk(), belum ditulis).

    Sumber data (semua sudah dihitung sebelumnya, fungsi ini TIDAK
    menghitung ulang saldo apa pun):
        neraca_now/lalu    : output susun_neraca() (laporan_keuangan.py)
                              utk tanggal_now/tanggal_lalu (2x panggilan
                              terpisah, dilakukan PEMANGGIL)
        laba_rugi_now/lalu : output susun_laba_rugi() utk periode
                              berjalan & pembanding
        profil             : dict Note 1 "Umum" (lihat CONTOH_PROFIL)

    Args:
        aset_tetap: opsional, {"daftar_kategori", "nilai_buku_now",
            "nilai_buku_lalu", "beban_penyusutan_tahun_berjalan"} --
            lihat tulis_note_7_aset_tetap(). Kalau None/daftar_kategori
            kosong, Note Aset Tetap DI-SKIP -- bukan sumber data COA
            biasa (lihat catatan [ASET TETAP] di calk_mapping.py),
            pemanggil WAJIB isi manual dari modul mutasi aset tetap yg
            sudah ada di sistem.
        hasil_pph_badan: opsional, output PERSIS
            modules.pph_badan.hitung_pph_pasal_31e() -- dipakai kalau
            client skema Tarif Umum Pasal 17/31E.
        pph_final_umkm: opsional, {"peredaran_bruto_now",
            "pph_final_now", "peredaran_bruto_lalu", "pph_final_lalu",
            "tarif"} -- dipakai kalau client skema PPh Final UMKM (PP
            55/2022). Isi HANYA SALAH SATU sesuai skema_pajak client
            (cek modules/pph_badan.py) -- kalau keduanya diisi,
            hasil_pph_badan diprioritaskan. Kalau keduanya None, Note
            Perpajakan DI-SKIP (masuk ke "peringatan" hasil).
        pihak_berelasi / peristiwa_setelah_neraca: opsional,
            {"id": [...], "en": [...]} kalimat custom 1:1 index --
            kalau None, dipakai kalimat standar "tidak ada transaksi/
            peristiwa material yang perlu diungkapkan".
        tanggal_persetujuan: opsional -- default tanggal_now kalau
            tidak diisi (dicatat di "peringatan", BUKAN error, supaya
            generate tetap jalan -- konsisten kebijakan Fase 4 poin 12
            roadmap: "kalau data belum lengkap, tetap generate dengan
            placeholder, jangan gagal total").
        grouping_piutang_usaha: opsional -- {"cabang": [{"label_id",
            "label_en", "no_akun": [...]}, ...]}, dipetakan by no_akun
            (lihat calk_mapping.susun_grouping_piutang_usaha() utk
            format & contoh persis). Kalau diisi, Note 4 Piutang Usaha
            ditampilkan dgn subgrup cabang+channel (format PERSIS file
            referensi PT AADL) alih-alih flat. Kalau None (default),
            Note 4 tetap flat seperti sebelumnya -- TIDAK ADA PERUBAHAN
            PERILAKU dari versi sebelum fitur ini ditambahkan.

    Returns:
        {
          "nomor_note_terakhir": int,
          "daftar_note_ditulis": [{"nomor", "key", "judul_id"}, ...],
          "peringatan": [str, ...],  -- gabungan
              peringatan_sub_kategori_tidak_dikenal dari calk_mapping +
              peringatan lain (note yg di-skip, tanggal_persetujuan
              fallback, dst) -- TIDAK PERNAH memblokir generate, cuma
              penanda supaya akuntan bisa cek ulang.
        }
    """
    peringatan: List[str] = []
    daftar_note_ditulis: List[Dict[str, str]] = []

    # [FIX -- POINT 3/4 gaya] Cache terjemahan dimuat SEKALI di sini untuk
    # SELURUH dokumen CALK (bukan per note/12x) dan dioper ke tiap
    # panggilan note yang butuh terjemahan -- lihat catatan panjang di
    # tulis_note_akun_generik() & terjemahkan_id_ke_en(). Ditulis balik
    # ke disk SEKALI SAJA di akhir fungsi ini (bukan per terjemahan),
    # via _simpan_cache_terjemahan() sebelum return.
    cache_terjemahan: Dict[str, str] = _muat_cache_terjemahan()

    # --- Header dokumen (kop halaman pertama, dipanggil SEKALI) ---
    nama_perusahaan = _isi(profil, "nama_perusahaan")
    _tambah_header_dokumen(doc, nama_perusahaan, tanggal_now, tanggal_lalu)

    # --- Note 1 & 2: SELALU ada, tidak lewat mapping COA (Fase 1) ---
    tulis_note_1_umum(doc, "1", profil, tanggal_now=tanggal_now)
    tulis_note_2_kebijakan_akuntansi(doc, "2", profil)
    daftar_note_ditulis.append({"nomor": "1", "key": "umum", "judul_id": "UMUM"})
    daftar_note_ditulis.append({"nomor": "2", "key": "kebijakan_akuntansi",
                                 "judul_id": "IKHTISAR KEBIJAKAN AKUNTANSI"})
    nomor = 3

    # --- Note 3 dst: hasil mapping COA (Fase 2) ---
    hasil_mapping = kelompokkan_akun_calk(
        neraca_now, neraca_lalu, laba_rugi_now, laba_rugi_lalu,
        grouping_piutang_usaha=grouping_piutang_usaha,
    )
    peringatan.extend(
        f'Akun "{p["nama_akun"]}" ({p["no_akun"]}) sub_kategori tidak dikenal '
        f'("{p["sub_kategori"]}"), otomatis masuk note "{p["masuk_ke"]}" -- cek COA.'
        for p in hasil_mapping["peringatan_sub_kategori_tidak_dikenal"]
    )
    peringatan.extend(hasil_mapping.get("peringatan_grouping_piutang_usaha") or [])

    for key, cfg in hasil_mapping["notes"].items():
        if key == "aset_tetap":
            if not aset_tetap or not aset_tetap.get("daftar_kategori"):
                continue  # tidak disuplai pemanggil -- skip, lihat docstring
            tulis_note_7_aset_tetap(
                doc, str(nomor), aset_tetap["daftar_kategori"],
                aset_tetap.get("nilai_buku_now", 0.0), aset_tetap.get("nilai_buku_lalu", 0.0),
                tanggal_now, tanggal_lalu,
                beban_penyusutan_tahun_berjalan=aset_tetap.get("beban_penyusutan_tahun_berjalan"),
            )
            daftar_note_ditulis.append({"nomor": str(nomor), "key": key, "judul_id": cfg["judul_id"]})
            nomor += 1
            continue

        if cfg["kosong"]:
            continue  # tidak ada akun -- note tidak ditampilkan (roadmap Fase 3 poin 7)

        if key == "kas":
            kas_bank, deposito = _pisah_kas_bank_dan_deposito(cfg["daftar_akun"])
            tulis_note_3_kas_dan_setara_kas(doc, str(nomor), kas_bank, deposito,
                                             tanggal_now, tanggal_lalu,
                                             cache_terjemahan=cache_terjemahan)
        elif cfg["fungsi"] == "akun_generik":
            label_j_id = cfg.get("label_jumlah_id") or f'Jumlah {cfg["judul_id"].title()}'
            label_j_en = cfg.get("label_jumlah_en") or f'Total {cfg["judul_en"].title()}'
            tulis_note_akun_generik(
                doc, str(nomor), cfg["judul_id"], cfg["judul_en"], cfg["daftar_akun"],
                tanggal_now, tanggal_lalu,
                kalimat_pembuka_id=f'Rincian {cfg["judul_id"].lower()} adalah sebagai berikut:',
                kalimat_pembuka_en=f'The details of {cfg["judul_en"].lower()} are as follows:',
                label_jumlah_id=label_j_id, label_jumlah_en=label_j_en,
                cache_terjemahan=cache_terjemahan,
            )
        else:
            fungsi = _FUNGSI_NOTE_STANDAR.get(cfg["fungsi"])
            if fungsi is None:
                peringatan.append(
                    f'Note "{key}" ({cfg["fungsi"]}) belum terdaftar di '
                    f'_FUNGSI_NOTE_STANDAR orchestrator -- note DILEWATI, cek '
                    f'calk_mapping.DAFTAR_NOTE_CALK vs calk_export.py.'
                )
                continue
            fungsi(doc, str(nomor), cfg["daftar_akun"], tanggal_now, tanggal_lalu,
                   cache_terjemahan=cache_terjemahan)

        daftar_note_ditulis.append({"nomor": str(nomor), "key": key, "judul_id": cfg["judul_id"]})
        nomor += 1

    # --- Note Perpajakan: sumber khusus (pph_badan.py), bukan mapping COA ---
    if hasil_pph_badan:
        tulis_note_15_perpajakan(doc, str(nomor), hasil_pph_badan, tanggal_now, tanggal_lalu,
                                  pph_badan_terutang_lalu=pph_badan_terutang_lalu)
        daftar_note_ditulis.append({"nomor": str(nomor), "key": "perpajakan", "judul_id": "PERPAJAKAN"})
        nomor += 1
    elif pph_final_umkm:
        tulis_note_15_perpajakan_final_umkm(
            doc, str(nomor),
            pph_final_umkm["peredaran_bruto_now"], pph_final_umkm["pph_final_now"],
            pph_final_umkm.get("peredaran_bruto_lalu"), pph_final_umkm.get("pph_final_lalu"),
            tanggal_now, tanggal_lalu, tarif=pph_final_umkm.get("tarif", 0.005),
        )
        daftar_note_ditulis.append({"nomor": str(nomor), "key": "perpajakan", "judul_id": "PERPAJAKAN"})
        nomor += 1
    else:
        peringatan.append(
            "Note Perpajakan DI-SKIP -- hasil_pph_badan/pph_final_umkm tidak "
            "diisi pemanggil. Cek skema_pajak client (modules/pph_badan.py)."
        )

    # --- Note penutup: SELALU ada, urutan tetap (standar SAK ETAP/EMKM) ---
    tulis_note_pihak_berelasi(
        doc, str(nomor),
        daftar_transaksi_id=(pihak_berelasi or {}).get("id"),
        daftar_transaksi_en=(pihak_berelasi or {}).get("en"),
    )
    daftar_note_ditulis.append({"nomor": str(nomor), "key": "pihak_berelasi",
                                 "judul_id": "TRANSAKSI DENGAN PIHAK-PIHAK BERELASI"})
    nomor += 1

    tulis_note_peristiwa_setelah_tanggal_neraca(
        doc, str(nomor), tanggal_now,
        daftar_peristiwa_id=(peristiwa_setelah_neraca or {}).get("id"),
        daftar_peristiwa_en=(peristiwa_setelah_neraca or {}).get("en"),
    )
    daftar_note_ditulis.append({"nomor": str(nomor), "key": "peristiwa_setelah_neraca",
                                 "judul_id": "PERISTIWA SETELAH TANGGAL NERACA"})
    nomor += 1

    if tanggal_persetujuan is None:
        tanggal_persetujuan = tanggal_now
        peringatan.append(
            "tanggal_persetujuan tidak diisi pemanggil -- dipakai tanggal_now "
            f"({_tgl_id(tanggal_now)}) sbg fallback, WAJIB dikonfirmasi ulang "
            "ke akuntan/direksi sebelum laporan final diterbitkan."
        )
    tulis_note_persetujuan_laporan_keuangan(
        doc, str(nomor), tanggal_persetujuan,
        nama_penanggung_jawab_id=nama_penanggung_jawab_id,
        nama_penanggung_jawab_en=nama_penanggung_jawab_en,
    )
    daftar_note_ditulis.append({"nomor": str(nomor), "key": "persetujuan_laporan_keuangan",
                                 "judul_id": "PERSETUJUAN DAN PENERBITAN LAPORAN KEUANGAN"})

    # [FIX -- POINT 3/4 gaya] Flush cache terjemahan SEKALI di sini,
    # setelah SEMUA note (dan semua terjemahan barunya) selesai diproses
    # -- menggantikan pola lama (flush per terjemahan, di dalam
    # terjemahkan_id_ke_en()). Kalau tidak ada terjemahan baru sama
    # sekali (semua sudah ada di cache / tidak ada akun berketerangan),
    # ini masih aman ditulis (isi sama dengan yang dibaca di awal).
    _simpan_cache_terjemahan(cache_terjemahan)

    return {
        "nomor_note_terakhir": nomor,
        "daftar_note_ditulis": daftar_note_ditulis,
        "peringatan": peringatan,
    }

# ============================================================
# 22. ENTRY POINT [FASE 4 -- roadmap CALK] -- susun semua note ->
#     doc.save() -> convert PDF (LibreOffice headless)
# ============================================================

def export_calk(
    output_dir: str,
    nama_file_dasar: str,
    profil: Dict[str, Any],
    neraca_now: Dict[str, Any], neraca_lalu: Dict[str, Any],
    laba_rugi_now: Dict[str, Any], laba_rugi_lalu: Dict[str, Any],
    tanggal_now: date, tanggal_lalu: date,
    **kwargs_orchestrator: Any,
) -> Dict[str, Any]:
    """
    [FASE 4] Entry point akhir: bikin Document() baru -> susun SEMUA
    note (Fase 3, susun_dan_tulis_semua_note_calk()) -> doc.save()
    -> convert ke PDF (LibreOffice headless, pola SAMA dgn yang sudah
    dipakai & terverifikasi jalan di scripts/recalc.py -- TIDAK
    reinvent) -> return path kedua file + hasil orchestrator (nomor
    note terakhir, daftar note yang ditulis, peringatan QA).

    kwargs_orchestrator diteruskan APA ADANYA ke
    susun_dan_tulis_semua_note_calk() -- aset_tetap, hasil_pph_badan,
    pph_final_umkm, pihak_berelasi, peristiwa_setelah_neraca,
    tanggal_persetujuan, dst (lihat docstring fungsi itu).

    [Fase 4 poin 12 roadmap -- "kalau data belum lengkap, tetap
    generate, jangan gagal total"] Fungsi ini TIDAK memvalidasi
    kelengkapan `profil` -- field kosong otomatis jadi placeholder
    "-- lengkapi data --" lewat _isi() di tulis_note_1_umum(), generate
    tetap jalan. Error yang DITANGKAP & di-raise ulang dgn pesan lebih
    jelas di sini HANYA error TEKNIS (soffice tidak ada / gagal
    convert / permission folder) -- BUKAN dibungkam, karena kalau file
    output rusak/tidak lengkap, pemanggil (endpoint Fase 5) wajib tahu
    supaya tidak menyajikan file rusak ke user.

    Returns:
        {"docx": path .docx, "pdf": path .pdf,
         "hasil_orchestrator": {"nomor_note_terakhir", "daftar_note_ditulis",
                                 "peringatan"} -- return value
             susun_dan_tulis_semua_note_calk() apa adanya, supaya
             pemanggil (endpoint) bisa log/tampilkan peringatan QA ke
             akuntan tanpa generate ulang.}
    """
    os.makedirs(output_dir, exist_ok=True)

    # [FIX -- KEAMANAN] nama_file_dasar bisa berasal dari input user tanpa
    # validasi pola (mis. main.py membangunnya dari req.periode_now, field
    # bertipe str bebas di CalkGenerateRequest) -- SEBELUM fix ini,
    # nama_file_dasar dipakai apa adanya di os.path.join(output_dir, ...)
    # di bawah, jadi nilai seperti "../../../tmp/x" atau path absolut bisa
    # membuat docx/pdf tertulis KELUAR dari output_dir (path traversal saat
    # MENULIS file, bukan cuma saat mengunduhnya). Endpoint download
    # (/api/unduh/{nama_file} di main.py) sudah menyaring dgn
    # Path(nama_file).name, tapi sisi tulis di sini belum -- disamakan di
    # sini: ambil basename saja & buang karakter yang bukan alfanumerik/
    # underscore/dash, supaya path_docx/path_pdf DIJAMIN tetap di dalam
    # output_dir apa pun isi nama_file_dasar yang dikirim pemanggil.
    nama_file_dasar = os.path.basename(nama_file_dasar or "CALK")
    nama_file_dasar = re.sub(r"[^A-Za-z0-9_\-]", "_", nama_file_dasar).strip("_") or "CALK"

    doc = Document()

    # Margin standar dokumen legal Indonesia -- TIDAK diatur di note
    # manapun sebelumnya (semua note cuma atur lebar tabel internal,
    # bukan margin halaman), jadi diset di sini sbg tanggung jawab
    # entry point.
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

    hasil_orch = susun_dan_tulis_semua_note_calk(
        doc, profil=profil,
        neraca_now=neraca_now, neraca_lalu=neraca_lalu,
        laba_rugi_now=laba_rugi_now, laba_rugi_lalu=laba_rugi_lalu,
        tanggal_now=tanggal_now, tanggal_lalu=tanggal_lalu,
        **kwargs_orchestrator,
    )
    # Footer nomor halaman ditambahkan di sini (bukan di orchestrator)
    # -- orchestrator (Fase 3) cuma urus ISI note, tata letak
    # halaman/dokumen (margin, footer) tetap tanggung jawab entry point
    # Fase 4, supaya orchestrator tetap bisa dipanggil terpisah (mis.
    # utk preview/testing) tanpa efek samping ke layout dokumen.
    _atur_footer_nomor_halaman(doc)

    if hasil_orch.get("peringatan"):
        for pesan in hasil_orch["peringatan"]:
            logger.warning(f"CALK [{nama_file_dasar}]: {pesan}")

    path_docx = os.path.join(output_dir, f"{nama_file_dasar}.docx")
    doc.save(path_docx)
    logger.info(f"CALK docx tersimpan: {path_docx}")

    path_pdf = _convert_docx_ke_pdf(path_docx, output_dir)
    logger.info(f"CALK pdf tersimpan: {path_pdf}")

    return {"docx": path_docx, "pdf": path_pdf, "hasil_orchestrator": hasil_orch}


def _convert_docx_ke_pdf(path_docx: str, output_dir: str) -> str:
    """Convert .docx -> .pdf via LibreOffice headless -- pola SAMA
    dgn yang sudah dipakai & terverifikasi jalan di scripts/recalc.py
    (subprocess.run(["soffice","--headless",...])), TIDAK reinvent.
    Error DIBIARKAN naik (raise), bukan ditangkap diam-diam -- lihat
    catatan penanganan error di export_calk().

    [FIX -- RELIABILITAS] soffice --headless SECARA DEFAULT memakai satu
    "user installation" (profil) yang sama untuk semua proses di mesin
    yang sama. Kalau dua CALK digenerate BERSAMAAN (mis. 2 supervisor beda
    client, atau 1 request retry sementara request sebelumnya belum
    selesai), proses soffice kedua bisa gagal/hang karena rebutan lock
    profil itu -- ini bug klasik LibreOffice headless di server multi-user,
    BUKAN skenario langka untuk aplikasi ini (endpoint /calk/generate bisa
    dipanggil kapan saja oleh Supervisor mana pun). Fix: pakai
    "-env:UserInstallation" unik per panggilan (folder temp berdasarkan
    nama file docx) supaya proses konversi paralel tidak saling kunci.
    """
    import tempfile
    import uuid

    profil_temp = os.path.join(
        tempfile.gettempdir(), f"lo_profile_calk_{uuid.uuid4().hex}"
    )
    try:
        hasil = subprocess.run(
            ["soffice",
             f"-env:UserInstallation=file://{profil_temp}",
             "--headless", "--convert-to", "pdf",
             "--outdir", output_dir, path_docx],
            capture_output=True, text=True, timeout=120, check=True,
        )
        if hasil.stdout.strip():
            logger.info(f"soffice convert stdout: {hasil.stdout.strip()}")
    except FileNotFoundError as e:
        raise RuntimeError(
            "LibreOffice ('soffice') tidak ditemukan di server -- CALK "
            "docx tetap tersimpan, tapi convert PDF gagal. Cek instalasi "
            "LibreOffice di server (lihat scripts/recalc.py utk pola yang "
            "sudah jalan di fitur lain)."
        ) from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            f"soffice gagal convert {os.path.basename(path_docx)} ke PDF "
            f"(exit code {e.returncode}): {e.stderr.strip()}"
        ) from e
    except subprocess.TimeoutExpired as e:
        raise RuntimeError(
            f"soffice convert timeout (>120 detik) utk "
            f"{os.path.basename(path_docx)} -- dokumen mungkin terlalu "
            f"besar/kompleks, atau server sedang berat."
        ) from e
    finally:
        # Profil temp cuma dipakai sekali per konversi -- bersihkan supaya
        # /tmp tidak menumpuk folder profil LibreOffice tiap CALK digenerate.
        # Kegagalan hapus (mis. masih terkunci sesaat) TIDAK menggagalkan
        # konversi yang sudah selesai -- cukup dicatat sbg warning.
        import shutil
        shutil.rmtree(profil_temp, ignore_errors=True)

    nama_dasar = os.path.splitext(os.path.basename(path_docx))[0]
    path_pdf = os.path.join(output_dir, f"{nama_dasar}.pdf")
    if not os.path.exists(path_pdf):
        raise RuntimeError(
            f"soffice sukses tanpa error tapi file PDF tidak ditemukan di "
            f"{path_pdf} -- cek permission folder output_dir."
        )
    return path_pdf